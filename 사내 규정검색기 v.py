# -*- coding: utf-8 -*-
"""
사내 규정 검색 프로그램 v8.1 (로컬 AI 기반)
EXE 패키징 최적화 버전 - 경량화, 지연 임포트, 메모리 최적화
"""

from __future__ import annotations

import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
import os
import sys
import json
import tempfile
import hashlib
import shutil
import logging
import subprocess
import platform
import re
import gc
from typing import List, Dict, Tuple, Optional, Callable, Any, TYPE_CHECKING
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum, auto
from collections import Counter
import math

# TYPE_CHECKING으로 타입 힌트만 사용 (런타임에 임포트 안 함)
if TYPE_CHECKING:
    from langchain_community.vectorstores import FAISS
    from langchain_huggingface import HuggingFaceEmbeddings

# ============================================================================
# 상수 및 설정
# ============================================================================
class AppConfig:
    """애플리케이션 설정"""
    APP_NAME = "사내 규정 검색기"
    APP_VERSION = "9.0"
    
    AVAILABLE_MODELS: Dict[str, str] = {
        "SNU SBERT (고성능)": "snunlp/KR-SBERT-V40K-klueNLI-augSTS",
        "BM-K Simal (균형)": "BM-K/ko-simal-roberta-base", 
        "JHGan SBERT (빠름)": "jhgan/ko-sbert-nli"
    }
    DEFAULT_MODEL = "JHGan SBERT (빠름)"  # 경량 모델 기본값
    
    CONFIG_FILE = "config.json"
    HISTORY_FILE = "search_history.json"
    SUPPORTED_EXTENSIONS = ('.txt', '.docx', '.pdf')
    
    MAX_FONT_SIZE = 32
    MIN_FONT_SIZE = 8
    DEFAULT_FONT_SIZE = 14
    DEFAULT_SEARCH_RESULTS = 3
    MAX_SEARCH_RESULTS = 10
    MAX_HISTORY_SIZE = 30
    
    CHUNK_SIZE = 800  # 약간 줄여서 메모리 절약
    CHUNK_OVERLAP = 80
    
    VECTOR_WEIGHT = 0.7
    BM25_WEIGHT = 0.3
    
    # 메모리 최적화
    MAX_DOCS_IN_MEMORY = 5000  # 최대 문서 청크 수
    BATCH_SIZE = 100  # 배치 처리 크기


class TaskStatus(Enum):
    IDLE = auto()
    LOADING_MODEL = auto()
    PROCESSING_DOCS = auto()
    SEARCHING = auto()


class FileStatus(Enum):
    PENDING = "대기"
    PROCESSING = "처리중"
    SUCCESS = "완료"
    FAILED = "실패"
    CACHED = "캐시"


@dataclass
class TaskResult:
    success: bool
    message: str
    data: Any = None
    failed_items: List[str] = field(default_factory=list)


@dataclass 
class FileInfo:
    path: str
    name: str
    extension: str
    size: int
    status: FileStatus = FileStatus.PENDING
    chunks: int = 0
    error: str = ""


# ============================================================================
# 유틸리티
# ============================================================================
def get_app_directory() -> str:
    """앱 디렉토리 반환"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def get_resource_path(relative_path: str) -> str:
    """리소스 경로 반환 (PyInstaller 호환)"""
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


# 로거 설정 (경량화)
def setup_logger() -> logging.Logger:
    logger = logging.getLogger('RegSearch')
    if logger.handlers:
        return logger
    
    logger.setLevel(logging.INFO)
    
    # 파일 핸들러
    log_dir = os.path.join(get_app_directory(), 'logs')
    os.makedirs(log_dir, exist_ok=True)
    
    fh = logging.FileHandler(
        os.path.join(log_dir, f'app_{datetime.now():%Y%m%d}.log'),
        encoding='utf-8'
    )
    fh.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(fh)
    
    # 콘솔 (개발 시에만)
    if not getattr(sys, 'frozen', False):
        ch = logging.StreamHandler()
        ch.setLevel(logging.INFO)
        logger.addHandler(ch)
    
    return logger

logger = setup_logger()


class FileUtils:
    """파일 유틸리티"""
    
    @staticmethod
    def safe_read(path: str, encoding: str = 'utf-8') -> Tuple[Optional[str], Optional[str]]:
        try:
            with open(path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read(), None
        except Exception as e:
            return None, str(e)
    
    @staticmethod
    def get_metadata(path: str) -> Optional[Dict]:
        try:
            stat = os.stat(path)
            return {'size': stat.st_size, 'mtime': stat.st_mtime}
        except OSError as e:
            logger.debug(f"파일 메타데이터 조회 실패: {path} - {e}")
            return None
    
    @staticmethod
    def open_file(path: str):
        try:
            if platform.system() == 'Windows':
                os.startfile(path)
            elif platform.system() == 'Darwin':
                subprocess.run(['open', path], check=False)
            else:
                subprocess.run(['xdg-open', path], check=False)
        except Exception as e:
            logger.error(f"파일 열기 실패: {e}")
    
    @staticmethod
    def format_size(size: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f}{unit}"
            size /= 1024
        return f"{size:.1f}TB"


# ============================================================================
# BM25 (경량 구현)
# ============================================================================
class BM25Light:
    """경량 BM25 구현"""
    
    __slots__ = ['k1', 'b', 'corpus', 'doc_lens', 'avgdl', 'idf', 'N']
    
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus: List[List[str]] = []
        self.doc_lens: List[int] = []
        self.avgdl = 0.0
        self.idf: Dict[str, float] = {}
        self.N = 0
    
    def _tokenize(self, text: str) -> List[str]:
        """한국어 개선 토크나이징"""
        if not text:
            return []
        text = re.sub(r'[^\w\s가-힣]', ' ', text.lower())
        # 한글 조사/어미 제거를 위한 기본 처리
        tokens = [t for t in text.split() if len(t) >= 2]
        return tokens
    
    def fit(self, docs: List[str]):
        self.corpus = []
        self.doc_lens = []
        df = Counter()
        
        for doc in docs:
            tokens = self._tokenize(doc)
            self.corpus.append(tokens)
            self.doc_lens.append(len(tokens))
            df.update(set(tokens))
        
        self.N = len(docs)
        self.avgdl = sum(self.doc_lens) / self.N if self.N else 0
        
        # IDF 계산
        self.idf = {
            t: math.log((self.N - f + 0.5) / (f + 0.5) + 1)
            for t, f in df.items()
        }
        
        # 메모리 정리
        del df
        gc.collect()
    
    def search(self, query: str, top_k: int = 5) -> List[Tuple[int, float]]:
        """BM25 검색 수행"""
        if not self.corpus or not query:
            return []
        
        q_tokens = self._tokenize(query)
        if not q_tokens:
            return []
        
        scores = []
        
        for idx, doc_tokens in enumerate(self.corpus):
            if not doc_tokens:
                continue
            score = self._score(q_tokens, doc_tokens, self.doc_lens[idx])
            if score > 0:
                scores.append((idx, score))
        
        scores.sort(key=lambda x: x[1], reverse=True)
        return scores[:top_k]
    
    def _score(self, query: List[str], doc: List[str], doc_len: int) -> float:
        score = 0.0
        doc_tf = Counter(doc)
        
        for term in query:
            if term not in self.idf:
                continue
            tf = doc_tf.get(term, 0)
            idf = self.idf[term]
            
            num = tf * (self.k1 + 1)
            den = tf + self.k1 * (1 - self.b + self.b * doc_len / self.avgdl)
            score += idf * num / den if den > 0 else 0
        
        return score
    
    def clear(self):
        """메모리 해제"""
        self.corpus.clear()
        self.doc_lens.clear()
        self.idf.clear()
        gc.collect()


# ============================================================================
# 문서 추출기 (지연 임포트)
# ============================================================================
class DocumentExtractor:
    """문서 추출기 - 지연 임포트로 메모리 최적화"""
    
    def __init__(self):
        self._docx_module = None
        self._pdf_module = None
    
    @property
    def docx(self):
        if self._docx_module is None:
            try:
                from docx import Document
                self._docx_module = Document
            except ImportError:
                self._docx_module = False
        return self._docx_module
    
    @property
    def pdf(self):
        if self._pdf_module is None:
            try:
                from pypdf import PdfReader
                self._pdf_module = PdfReader
            except ImportError:
                self._pdf_module = False
        return self._pdf_module
    
    def extract(self, path: str) -> Tuple[str, Optional[str]]:
        """파일에서 텍스트 추출"""
        if not path or not os.path.exists(path):
            return "", f"파일 없음: {path}"
        
        if not os.path.isfile(path):
            return "", f"파일이 아님: {path}"
        
        ext = os.path.splitext(path)[1].lower()
        
        if ext == '.txt':
            return self._extract_txt(path)
        elif ext == '.docx':
            return self._extract_docx(path)
        elif ext == '.pdf':
            return self._extract_pdf(path)
        return "", f"지원하지 않는 형식: {ext}"
    
    def _extract_txt(self, path: str) -> Tuple[str, Optional[str]]:
        return FileUtils.safe_read(path)
    
    def _extract_docx(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.docx:
            return "", "DOCX 라이브러리 없음"
        
        try:
            doc = self.docx(path)
            parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            
            return '\n\n'.join(parts), None
        except Exception as e:
            return "", f"DOCX 오류: {e}"
    
    def _extract_pdf(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.pdf:
            return "", "PDF 라이브러리 없음"
        
        try:
            reader = self.pdf(path)
            
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception as e:
                    logger.warning(f"PDF 복호화 실패: {e}")
                    return "", "암호화된 PDF"
            
            texts = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
                except Exception as e:
                    logger.debug(f"PDF 페이지 추출 실패: {e}")
                    continue
            
            if not texts:
                return "", "텍스트 없음 (이미지 PDF)"
            
            return '\n\n'.join(texts), None
        except Exception as e:
            return "", f"PDF 오류: {e}"


# ============================================================================
# QA 시스템 (지연 임포트 + 메모리 최적화)
# ============================================================================
class RegulationQASystem:
    """규정 Q&A 시스템 - 경량화 버전"""
    
    def __init__(self):
        self.vector_store: Optional[FAISS] = None
        self.embedding_model: Optional[HuggingFaceEmbeddings] = None
        self.model_id: Optional[str] = None
        self.extractor = DocumentExtractor()
        self.cache_path = os.path.join(tempfile.gettempdir(), "reg_qa_v90")
        
        self.bm25: Optional[BM25Light] = None
        self.documents: List[str] = []
        self.doc_meta: List[Dict] = []
        self.file_infos: Dict[str, FileInfo] = {}
        self.current_folder = ""
        
        self._lock = threading.Lock()
    
    def load_model(
        self, 
        model_name: str,
        progress_cb: Optional[Callable[[str], None]] = None
    ) -> TaskResult:
        """모델 로드 (지연 임포트)"""
        model_id = AppConfig.AVAILABLE_MODELS.get(
            model_name, 
            AppConfig.AVAILABLE_MODELS[AppConfig.DEFAULT_MODEL]
        )
        
        try:
            if progress_cb:
                progress_cb("라이브러리 로드 중...")
            
            # 지연 임포트
            import torch
            from langchain_huggingface import HuggingFaceEmbeddings
            
            if progress_cb:
                progress_cb("모델 로딩 중...")
            
            cache_dir = os.path.join(get_app_directory(), 'models')
            os.makedirs(cache_dir, exist_ok=True)
            
            device = 'cuda' if torch.cuda.is_available() else 'cpu'
            
            self.embedding_model = HuggingFaceEmbeddings(
                model_name=model_id,
                cache_folder=cache_dir,
                model_kwargs={'device': device},
                encode_kwargs={'normalize_embeddings': True}
            )
            
            self.model_id = model_id
            
            # 메모리 정리
            gc.collect()
            if device == 'cuda':
                torch.cuda.empty_cache()
            
            logger.info(f"모델 로드 완료: {model_name} ({device})")
            return TaskResult(True, f"모델 로드 완료 ({device})")
            
        except Exception as e:
            logger.error(f"모델 로드 실패: {e}")
            return TaskResult(False, f"모델 로드 실패: {e}")
    
    def _get_cache_dir(self, folder: str) -> str:
        h1 = hashlib.md5(self.model_id.encode()).hexdigest()[:6]
        h2 = hashlib.md5(folder.encode()).hexdigest()[:6]
        return os.path.join(self.cache_path, f"{h2}_{h1}")
    
    def process_documents(
        self,
        folder: str,
        files: List[str],
        progress_cb: Callable[[int, str], None]
    ) -> TaskResult:
        """문서 처리 (증분 인덱싱)"""
        if not self.embedding_model:
            return TaskResult(False, "모델이 로드되지 않았습니다")
        
        with self._lock:
            return self._process_internal(folder, files, progress_cb)
    
    def _process_internal(
        self,
        folder: str,
        files: List[str],
        progress_cb: Callable[[int, str], None]
    ) -> TaskResult:
        # 지연 임포트
        from langchain.text_splitter import CharacterTextSplitter
        from langchain_community.vectorstores import FAISS
        from langchain.docstore.document import Document
        
        self.current_folder = folder
        cache_dir = self._get_cache_dir(folder)
        
        # 파일 정보 초기화
        self.file_infos.clear()
        for fp in files:
            meta = FileUtils.get_metadata(fp)
            self.file_infos[fp] = FileInfo(
                path=fp,
                name=os.path.basename(fp),
                extension=os.path.splitext(fp)[1].lower(),
                size=meta['size'] if meta else 0
            )
        
        # 캐시 정보 로드
        progress_cb(5, "캐시 확인...")
        cache_info = self._load_cache_info(cache_dir)
        
        # 변경 파일 확인
        to_process = []
        cached = []
        
        for fp in files:
            fname = os.path.basename(fp)
            meta = FileUtils.get_metadata(fp)
            
            if meta and fname in cache_info:
                cm = cache_info[fname]
                if cm.get('size') == meta['size'] and cm.get('mtime') == meta['mtime']:
                    cached.append(fp)
                    self.file_infos[fp].status = FileStatus.CACHED
                    self.file_infos[fp].chunks = cm.get('chunks', 0)
                    continue
            
            to_process.append(fp)
        
        # 캐시 로드
        self.documents = []
        self.doc_meta = []
        
        if cached and os.path.exists(os.path.join(cache_dir, "index.faiss")):
            try:
                progress_cb(10, "캐시 로드...")
                self.vector_store = FAISS.load_local(
                    cache_dir,
                    self.embedding_model,
                    allow_dangerous_deserialization=True
                )
                
                docs_path = os.path.join(cache_dir, "docs.json")
                if os.path.exists(docs_path):
                    with open(docs_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        self.documents = data.get('docs', [])
                        self.doc_meta = data.get('meta', [])
                
                logger.info(f"캐시 로드: {len(cached)}개 파일")
            except Exception as e:
                logger.warning(f"캐시 로드 실패: {e}")
                to_process = files
                cached = []
                self.vector_store = None
        
        # 처리할 파일 없음
        if not to_process:
            self._build_bm25()
            progress_cb(100, "완료!")
            return TaskResult(
                True,
                f"캐시에서 {len(cached)}개 파일 로드",
                {'chunks': len(self.documents), 'cached': len(cached), 'new': 0}
            )
        
        # 새 파일 처리
        splitter = CharacterTextSplitter(
            separator="\n\n",
            chunk_size=AppConfig.CHUNK_SIZE,
            chunk_overlap=AppConfig.CHUNK_OVERLAP
        )
        
        failed = []
        new_docs = []
        new_cache_info = {}
        
        total = len(to_process)
        for i, fp in enumerate(to_process):
            fname = os.path.basename(fp)
            pct = 15 + int((i / total) * 55)
            progress_cb(pct, f"처리: {fname}")
            
            self.file_infos[fp].status = FileStatus.PROCESSING
            
            try:
                content, error = self.extractor.extract(fp)
                
                if error:
                    failed.append(f"{fname} ({error})")
                    self.file_infos[fp].status = FileStatus.FAILED
                    self.file_infos[fp].error = error
                    continue
                
                if not content.strip():
                    failed.append(f"{fname} (빈 파일)")
                    self.file_infos[fp].status = FileStatus.FAILED
                    self.file_infos[fp].error = "빈 파일"
                    continue
                
                chunks = splitter.split_text(content)
                chunk_count = 0
                
                for chunk in chunks:
                    if chunk.strip():
                        new_docs.append(Document(
                            page_content=chunk.strip(),
                            metadata={"source": fname, "path": fp}
                        ))
                        self.documents.append(chunk.strip())
                        self.doc_meta.append({"source": fname, "path": fp})
                        chunk_count += 1
                
                self.file_infos[fp].status = FileStatus.SUCCESS
                self.file_infos[fp].chunks = chunk_count
                
                meta = FileUtils.get_metadata(fp)
                if meta:
                    new_cache_info[fname] = {
                        'size': meta['size'],
                        'mtime': meta['mtime'],
                        'chunks': chunk_count
                    }
                
                # 주기적 GC
                if (i + 1) % 10 == 0:
                    gc.collect()
                    
            except Exception as e:
                failed.append(f"{fname} ({e})")
                self.file_infos[fp].status = FileStatus.FAILED
                self.file_infos[fp].error = str(e)
        
        if not new_docs and not self.vector_store:
            return TaskResult(False, "처리 가능한 문서 없음", failed_items=failed)
        
        # 문서 수 제한
        if len(self.documents) > AppConfig.MAX_DOCS_IN_MEMORY:
            logger.warning(f"문서 수 제한 초과: {len(self.documents)}")
            self.documents = self.documents[:AppConfig.MAX_DOCS_IN_MEMORY]
            self.doc_meta = self.doc_meta[:AppConfig.MAX_DOCS_IN_MEMORY]
        
        # 벡터 저장소 업데이트
        progress_cb(75, "벡터 인덱스 생성...")
        
        try:
            if new_docs:
                if self.vector_store:
                    # 배치 추가
                    for i in range(0, len(new_docs), AppConfig.BATCH_SIZE):
                        batch = new_docs[i:i + AppConfig.BATCH_SIZE]
                        self.vector_store.add_documents(batch)
                else:
                    self.vector_store = FAISS.from_documents(
                        new_docs,
                        self.embedding_model
                    )
        except Exception as e:
            logger.error(f"벡터 인덱스 생성 실패: {e}")
            return TaskResult(False, f"인덱스 생성 실패: {e}")
        
        # BM25 빌드
        progress_cb(85, "키워드 인덱스 생성...")
        self._build_bm25()
        
        # 캐시 저장
        progress_cb(90, "캐시 저장...")
        self._save_cache(cache_dir, cache_info, new_cache_info)
        
        # 메모리 정리
        del new_docs
        gc.collect()
        
        progress_cb(100, "완료!")
        
        return TaskResult(
            True,
            f"{len(files) - len(failed)}개 처리 완료",
            {
                'chunks': len(self.documents),
                'new': len(to_process) - len(failed),
                'cached': len(cached)
            },
            failed
        )
    
    def _build_bm25(self):
        """BM25 인덱스 빌드"""
        if self.documents:
            self.bm25 = BM25Light()
            self.bm25.fit(self.documents)
        else:
            self.bm25 = None
    
    def _load_cache_info(self, cache_dir: str) -> Dict:
        path = os.path.join(cache_dir, "cache_info.json")
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError) as e:
                logger.warning(f"캐시 정보 로드 실패: {e}")
        return {}
    
    def _save_cache(self, cache_dir: str, old_info: Dict, new_info: Dict):
        try:
            os.makedirs(cache_dir, exist_ok=True)
            self.vector_store.save_local(cache_dir)
            
            # 캐시 정보 병합
            merged = {**old_info, **new_info}
            with open(os.path.join(cache_dir, "cache_info.json"), 'w', encoding='utf-8') as f:
                json.dump(merged, f, ensure_ascii=False)
            
            # 문서 저장
            with open(os.path.join(cache_dir, "docs.json"), 'w', encoding='utf-8') as f:
                json.dump({
                    'docs': self.documents,
                    'meta': self.doc_meta
                }, f, ensure_ascii=False)
                
        except Exception as e:
            logger.warning(f"캐시 저장 실패: {e}")
    
    def search(self, query: str, k: int = 3, hybrid: bool = True) -> TaskResult:
        """하이브리드 검색 - 개선된 버전"""
        if not self.vector_store:
            return TaskResult(False, "문서가 로드되지 않음")
        
        # 쿼리 검증
        query = query.strip() if query else ""
        if not query or len(query) < 2:
            return TaskResult(False, "검색어가 너무 짧습니다 (2자 이상)")
        
        try:
            k = max(1, min(k, AppConfig.MAX_SEARCH_RESULTS))
            
            # 벡터 검색
            vec_results = self.vector_store.similarity_search_with_score(query, k=k*2)
            
            results = {}
            
            if vec_results:
                # L2 거리 기반 점수 정규화: 거리가 작을수록 점수가 높음
                min_dist = min(r[1] for r in vec_results)
                max_dist = max(r[1] for r in vec_results)
                dist_range = max_dist - min_dist if max_dist != min_dist else 1
                
                for doc, dist in vec_results:
                    key = doc.page_content[:100]  # 키 길이 확장
                    # 정규화: 0~1 범위, 거리 작을수록 점수 높음
                    norm_score = 1 - ((dist - min_dist) / (dist_range + 0.001))
                    norm_score = max(0.1, min(1.0, norm_score))  # 최소 0.1 보장
                    
                    results[key] = {
                        'content': doc.page_content,
                        'source': doc.metadata.get('source', '?'),
                        'path': doc.metadata.get('path', ''),
                        'vec_score': norm_score,
                        'bm25_score': 0
                    }
            
            # BM25 검색
            if hybrid and self.bm25 and self.documents:
                bm25_res = self.bm25.search(query, top_k=k*2)
                
                if bm25_res:
                    max_bm = max(r[1] for r in bm25_res) if bm25_res else 1
                    for idx, score in bm25_res:
                        if 0 <= idx < len(self.documents):
                            content = self.documents[idx]
                            key = content[:100]
                            norm_score = score / (max_bm + 0.001) if max_bm > 0 else 0
                            
                            if key in results:
                                results[key]['bm25_score'] = norm_score
                            else:
                                meta = self.doc_meta[idx] if idx < len(self.doc_meta) else {}
                                results[key] = {
                                    'content': content,
                                    'source': meta.get('source', '?'),
                                    'path': meta.get('path', ''),
                                    'vec_score': 0,
                                    'bm25_score': norm_score
                                }
            
            # 최종 점수 계산
            for item in results.values():
                item['score'] = (
                    AppConfig.VECTOR_WEIGHT * item['vec_score'] +
                    AppConfig.BM25_WEIGHT * item['bm25_score']
                )
            
            sorted_res = sorted(results.values(), key=lambda x: x['score'], reverse=True)[:k]
            
            logger.debug(f"검색 완료: '{query}' -> {len(sorted_res)}개 결과")
            return TaskResult(True, "검색 완료", sorted_res)
            
        except Exception as e:
            logger.error(f"검색 오류: {e}", exc_info=True)
            return TaskResult(False, f"검색 오류: {e}")
    
    def get_file_infos(self) -> List[FileInfo]:
        return list(self.file_infos.values())
    
    def clear_cache(self) -> TaskResult:
        if os.path.exists(self.cache_path):
            try:
                shutil.rmtree(self.cache_path)
                return TaskResult(True, "캐시 삭제 완료")
            except Exception as e:
                return TaskResult(False, f"삭제 실패: {e}")
        return TaskResult(True, "삭제할 캐시 없음")
    
    def cleanup(self):
        """메모리 정리"""
        self.documents.clear()
        self.doc_meta.clear()
        if self.bm25:
            self.bm25.clear()
        gc.collect()


# ============================================================================
# 검색 히스토리
# ============================================================================
class SearchHistory:
    """검색 히스토리 관리"""
    
    def __init__(self):
        self.items: List[str] = []
        self.path = os.path.join(get_app_directory(), AppConfig.HISTORY_FILE)
        self._load()
    
    def _load(self):
        if os.path.exists(self.path):
            try:
                with open(self.path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.items = [h['q'] for h in data[:AppConfig.MAX_HISTORY_SIZE]]
            except (json.JSONDecodeError, IOError, KeyError) as e:
                logger.debug(f"히스토리 로드 실패: {e}")
                self.items = []
    
    def _save(self):
        try:
            with open(self.path, 'w', encoding='utf-8') as f:
                json.dump([{'q': q} for q in self.items], f, ensure_ascii=False)
        except IOError as e:
            logger.warning(f"히스토리 저장 실패: {e}")
    
    def add(self, query: str):
        self.items = [q for q in self.items if q != query]
        self.items.insert(0, query)
        self.items = self.items[:AppConfig.MAX_HISTORY_SIZE]
        self._save()
    
    def get(self, count: int = 10) -> List[str]:
        return self.items[:count]
    
    def clear(self):
        self.items = []
        self._save()


# ============================================================================
# UI 컴포넌트
# ============================================================================
class Toast(ctk.CTkToplevel):
    """토스트 알림 - 스택 지원"""
    _instances: List['Toast'] = []
    
    def __init__(self, parent, msg: str, duration: int = 2500, type_: str = "info"):
        super().__init__(parent)
        
        self.overrideredirect(True)
        self.attributes('-topmost', True)
        self.attributes('-alpha', 0.95)
        
        colors = {
            "info": ("#3B82F6", "ℹ️"),
            "success": ("#10B981", "✅"), 
            "warning": ("#F59E0B", "⚠️"),
            "error": ("#EF4444", "❌")
        }
        
        color, icon = colors.get(type_, ("#3B82F6", "ℹ️"))
        self.configure(fg_color=color, corner_radius=8)
        
        frame = ctk.CTkFrame(self, fg_color="transparent")
        frame.pack(fill="both", expand=True, padx=2, pady=2)
        
        ctk.CTkLabel(
            frame, text=f"{icon} {msg}", text_color="white",
            font=ctk.CTkFont(size=12, weight="bold"), wraplength=280
        ).pack(padx=14, pady=10)
        
        # 스택 위치 계산
        Toast._instances = [t for t in Toast._instances if t.winfo_exists()]
        Toast._instances.append(self)
        
        self.update_idletasks()
        offset = (len(Toast._instances) - 1) * 50
        x = parent.winfo_x() + parent.winfo_width() - self.winfo_width() - 20
        y = parent.winfo_y() + 55 + offset
        self.geometry(f"+{x}+{y}")
        
        self.after(duration, self._close)
    
    def _close(self):
        if self in Toast._instances:
            Toast._instances.remove(self)
        self.destroy()


class ProgressDlg(ctk.CTkToplevel):
    """진행 다이얼로그"""
    
    def __init__(self, parent, title: str = "처리 중"):
        super().__init__(parent)
        
        self.title(title)
        self.geometry("380x130")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()
        
        x = parent.winfo_x() + (parent.winfo_width() - 380) // 2
        y = parent.winfo_y() + (parent.winfo_height() - 130) // 2
        self.geometry(f"+{x}+{y}")
        
        self.status = ctk.CTkLabel(self, text="준비 중...", font=ctk.CTkFont(size=12))
        self.status.pack(pady=(25, 8))
        
        self.bar = ctk.CTkProgressBar(self, width=340)
        self.bar.pack(pady=8)
        self.bar.set(0)
        
        self.pct = ctk.CTkLabel(self, text="0%", font=ctk.CTkFont(size=11))
        self.pct.pack(pady=5)
        
        self.protocol("WM_DELETE_WINDOW", lambda: None)
    
    def update(self, percent: int, status: str = None):
        self.bar.set(percent / 100)
        self.pct.configure(text=f"{percent}%")
        if status:
            self.status.configure(text=status)
        super().update()


class HistoryPopup(ctk.CTkToplevel):
    """히스토리 팝업"""
    
    def __init__(self, parent, items: List[str], callback: Callable[[str], None]):
        super().__init__(parent)
        
        self.callback = callback
        self.overrideredirect(True)
        self.attributes('-topmost', True)
        
        frame = ctk.CTkFrame(self, corner_radius=6)
        frame.pack(fill="both", expand=True, padx=1, pady=1)
        
        for q in items[:8]:
            btn = ctk.CTkButton(
                frame, text=q, anchor="w",
                fg_color="transparent",
                hover_color=("gray80", "gray30"),
                height=28,
                command=lambda x=q: self._select(x)
            )
            btn.pack(fill="x", padx=4, pady=1)
        
        if not items:
            ctk.CTkLabel(frame, text="기록 없음", text_color="gray").pack(pady=8)
        
        self.bind("<FocusOut>", lambda e: self.destroy())
    
    def _select(self, q: str):
        self.callback(q)
        self.destroy()


class FileListPanel(ctk.CTkFrame):
    """파일 목록 패널"""
    
    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=10, pady=(10, 5))
        
        ctk.CTkLabel(
            header, text="📄 파일 목록",
            font=ctk.CTkFont(size=14, weight="bold")
        ).pack(side="left")
        
        self.count_lbl = ctk.CTkLabel(header, text="0개", text_color="gray")
        self.count_lbl.pack(side="right")
        
        self.scroll = ctk.CTkScrollableFrame(self)
        self.scroll.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.empty = ctk.CTkLabel(
            self.scroll, text="폴더를 선택하세요", text_color="gray"
        )
        self.empty.pack(pady=40)
    
    def update_files(self, infos: List[FileInfo]):
        for w in self.scroll.winfo_children():
            w.destroy()
        
        if not infos:
            ctk.CTkLabel(
                self.scroll, text="폴더를 선택하세요", text_color="gray"
            ).pack(pady=40)
            self.count_lbl.configure(text="0개")
            return
        
        self.count_lbl.configure(text=f"{len(infos)}개")
        
        # 통계 섹션
        stats_frame = ctk.CTkFrame(self.scroll)
        stats_frame.pack(fill="x", pady=(0, 8))
        
        total_chunks = sum(i.chunks for i in infos)
        total_size = sum(i.size for i in infos)
        success = sum(1 for i in infos if i.status in (FileStatus.SUCCESS, FileStatus.CACHED))
        failed = sum(1 for i in infos if i.status == FileStatus.FAILED)
        
        stats_text = f"📊 총 {len(infos)}개 파일 | {total_chunks} 청크 | {FileUtils.format_size(total_size)}"
        ctk.CTkLabel(
            stats_frame, text=stats_text,
            font=ctk.CTkFont(size=11, weight="bold")
        ).pack(pady=6, padx=10, anchor="w")
        
        status_text = f"✅ 성공: {success}  ❌ 실패: {failed}"
        ctk.CTkLabel(
            stats_frame, text=status_text,
            font=ctk.CTkFont(size=10), text_color="gray"
        ).pack(pady=(0, 6), padx=10, anchor="w")
        
        icons = {
            FileStatus.SUCCESS: "✅",
            FileStatus.CACHED: "💾",
            FileStatus.FAILED: "❌",
            FileStatus.PROCESSING: "⏳",
            FileStatus.PENDING: "⏸️"
        }
        
        for info in infos:
            row = ctk.CTkFrame(self.scroll, fg_color="transparent")
            row.pack(fill="x", pady=1)
            
            ctk.CTkLabel(row, text=icons.get(info.status, "?"), width=25).pack(side="left")
            
            ctk.CTkButton(
                row, text=info.name, anchor="w",
                fg_color="transparent",
                hover_color=("gray85", "gray25"),
                text_color=("black", "white"),
                height=26,
                command=lambda p=info.path: FileUtils.open_file(p)
            ).pack(side="left", fill="x", expand=True)
            
            if info.chunks:
                ctk.CTkLabel(row, text=f"{info.chunks}", text_color="gray", width=40).pack(side="right")
            
            ctk.CTkLabel(
                row, text=FileUtils.format_size(info.size),
        text_color="gray", width=60
            ).pack(side="right")


class ResultCard(ctk.CTkFrame):
    """검색 결과 카드 - 점수 시각화 및 하이라이트 지원"""
    
    def __init__(self, parent, idx: int, data: Dict, on_copy: Callable, 
                 font_size: int = 12, query: str = "", **kw):
        super().__init__(parent, **kw)
        
        score = data.get('score', 0)
        
        # 헤더
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", padx=10, pady=(8, 4))
        
        ctk.CTkLabel(
            hdr, text=f"▶ 결과 {idx}",
            font=ctk.CTkFont(size=12, weight="bold")
        ).pack(side="left")
        
        ctk.CTkLabel(
            hdr, text=data['source'],
            text_color="#3B82F6", font=ctk.CTkFont(size=11)
        ).pack(side="left", padx=(12, 0))
        
        # 점수 표시
        score_pct = min(100, int(score * 100))
        score_color = "#10B981" if score_pct >= 70 else "#F59E0B" if score_pct >= 40 else "#EF4444"
        
        score_frame = ctk.CTkFrame(hdr, fg_color="transparent")
        score_frame.pack(side="left", padx=(12, 0))
        
        ctk.CTkLabel(
            score_frame, text=f"{score_pct}%",
            font=ctk.CTkFont(size=10), text_color=score_color
        ).pack(side="left", padx=(0, 4))
        
        score_bar = ctk.CTkProgressBar(score_frame, width=50, height=8)
        score_bar.pack(side="left")
        score_bar.set(score)
        score_bar.configure(progress_color=score_color)
        
        # 버튼
        btn_f = ctk.CTkFrame(hdr, fg_color="transparent")
        btn_f.pack(side="right")
        
        ctk.CTkButton(
            btn_f, text="📋복사", width=60, height=24,
            command=lambda: on_copy(data['content'])
        ).pack(side="left", padx=2)
        
        if data.get('path'):
            ctk.CTkButton(
                btn_f, text="📂열기", width=60, height=24,
                command=lambda: FileUtils.open_file(data['path'])
            ).pack(side="left", padx=2)
        
        # 내용
        txt = ctk.CTkTextbox(self, height=100, wrap="word", font=("", font_size))
        txt.pack(fill="x", padx=10, pady=(4, 8))
        
        content = data['content']
        txt.insert("1.0", content)
        
        # 검색어 하이라이트
        if query:
            txt.tag_config("highlight", foreground="#F59E0B")
            for word in query.split():
                if len(word) >= 2:
                    start = "1.0"
                    while True:
                        pos = txt.search(word, start, stopindex="end", nocase=True)
                        if not pos:
                            break
                        end = f"{pos}+{len(word)}c"
                        txt.tag_add("highlight", pos, end)
                        start = end
        
        txt.configure(state="disabled")


# ============================================================================
# 메인 앱
# ============================================================================
class App(ctk.CTk):
    """메인 애플리케이션"""
    
    def __init__(self, qa: RegulationQASystem):
        super().__init__()
        
        self.qa = qa
        self.history = SearchHistory()
        self.chat_log: List[Dict] = []
        self.font_size = AppConfig.DEFAULT_FONT_SIZE
        self.last_folder = ""
        self.model_name = AppConfig.DEFAULT_MODEL
        self.status = TaskStatus.IDLE
        self.hybrid = True
        
        self._load_config()
        
        self.title(f"{AppConfig.APP_NAME} v{AppConfig.APP_VERSION}")
        self.geometry("1200x800")
        self.minsize(950, 650)
        
        self._set_icon()
        
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self._build_ui()
        self._bind_keys()
        
        self.after(150, self._init_app)
        self.protocol("WM_DELETE_WINDOW", self._on_close)
    
    def _set_icon(self):
        try:
            ico = os.path.join(get_app_directory(), 'icon.ico')
            if os.path.exists(ico):
                self.iconbitmap(ico)
        except Exception as e:
            logger.debug(f"아이콘 설정 실패: {e}")
    
    def _load_config(self):
        path = os.path.join(get_app_directory(), AppConfig.CONFIG_FILE)
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    cfg = json.load(f)
                self.last_folder = cfg.get("folder", "")
                self.model_name = cfg.get("model", AppConfig.DEFAULT_MODEL)
                self.font_size = cfg.get("font", AppConfig.DEFAULT_FONT_SIZE)
                self.hybrid = cfg.get("hybrid", True)
            except (json.JSONDecodeError, IOError) as e:
                logger.warning(f"설정 로드 실패: {e}")
    
    def _save_config(self):
        try:
            with open(os.path.join(get_app_directory(), AppConfig.CONFIG_FILE), 'w', encoding='utf-8') as f:
                json.dump({
                    "folder": self.last_folder,
                    "model": self.model_name,
                    "font": self.font_size,
                    "hybrid": self.hybrid
                }, f, ensure_ascii=False)
        except IOError as e:
            logger.warning(f"설정 저장 실패: {e}")
    
    def _bind_keys(self):
        self.bind("<Control-o>", lambda e: self._open_folder())
        self.bind("<Control-f>", lambda e: self.query_entry.focus())
        self.bind("<Control-h>", lambda e: self._show_history())
        self.bind("<Control-s>", lambda e: self._save_log())
        self.bind("<Escape>", lambda e: self.query_entry.delete(0, 'end'))
    
    # ========================================================================
    # UI 구축
    # ========================================================================
    def _build_ui(self):
        self.tabs = ctk.CTkTabview(self, anchor="w")
        self.tabs.grid(row=0, column=0, padx=12, pady=(8, 4), sticky="nsew")
        
        self.tabs.add("🔍 검색")
        self.tabs.add("📄 파일")
        self.tabs.add("⚙️ 설정")
        
        self._build_search_tab()
        self._build_files_tab()
        self._build_settings_tab()
        self._build_statusbar()
    
    def _build_search_tab(self):
        tab = self.tabs.tab("🔍 검색")
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(1, weight=1)
        
        # 컨트롤
        ctrl = ctk.CTkFrame(tab)
        ctrl.grid(row=0, column=0, padx=8, pady=8, sticky="ew")
        ctrl.grid_columnconfigure(2, weight=1)
        
        ctk.CTkLabel(
            ctrl, text="📁 규정 폴더",
            font=ctk.CTkFont(size=14, weight="bold")
        ).grid(row=0, column=0, columnspan=4, pady=(10, 12), padx=12, sticky="w")
        
        self.folder_btn = ctk.CTkButton(
            ctrl, text="📂 열기 (Ctrl+O)",
            command=self._open_folder, width=120, height=30, state="disabled"
        )
        self.folder_btn.grid(row=1, column=0, padx=(12, 4), pady=(0, 10))
        
        self.recent_btn = ctk.CTkButton(
            ctrl, text="🕐 최근",
            command=self._load_recent, width=80, height=30, state="disabled"
        )
        self.recent_btn.grid(row=1, column=1, padx=4, pady=(0, 10))
        
        self.folder_lbl = ctk.CTkLabel(ctrl, text="폴더를 선택하세요", text_color="gray", anchor="w")
        self.folder_lbl.grid(row=1, column=2, padx=10, sticky="ew", pady=(0, 10))
        
        self.refresh_btn = ctk.CTkButton(
            ctrl, text="🔄", width=35, height=30,
            command=self._refresh, state="disabled"
        )
        self.refresh_btn.grid(row=1, column=3, padx=(4, 12), pady=(0, 10))
        
        # 결과
        res_frame = ctk.CTkFrame(tab)
        res_frame.grid(row=1, column=0, padx=8, pady=4, sticky="nsew")
        res_frame.grid_rowconfigure(1, weight=1)
        res_frame.grid_columnconfigure(0, weight=1)
        
        # 툴바
        toolbar = ctk.CTkFrame(res_frame, fg_color="transparent")
        toolbar.grid(row=0, column=0, padx=8, pady=6, sticky="ew")
        
        ctk.CTkLabel(toolbar, text="글자:").pack(side="left", padx=(4, 6))
        ctk.CTkButton(toolbar, text="－", width=28, height=26, command=lambda: self._font(-2)).pack(side="left", padx=1)
        self.font_lbl = ctk.CTkLabel(toolbar, text=str(self.font_size), width=30)
        self.font_lbl.pack(side="left", padx=1)
        ctk.CTkButton(toolbar, text="＋", width=28, height=26, command=lambda: self._font(2)).pack(side="left", padx=(1, 12))
        
        ctk.CTkButton(toolbar, text="🔄 초기화", width=75, height=26, command=self._reset).pack(side="left", padx=4)
        ctk.CTkButton(toolbar, text="💾 저장", width=65, height=26, command=self._save_log).pack(side="right", padx=4)
        
        self.result_scroll = ctk.CTkScrollableFrame(res_frame)
        self.result_scroll.grid(row=1, column=0, sticky="nsew", padx=8, pady=(4, 8))
        
        ctk.CTkLabel(
            self.result_scroll, text="👋 폴더를 선택하고 검색하세요",
            text_color="gray", font=ctk.CTkFont(size=13)
        ).pack(pady=80)
        
        # 입력
        inp = ctk.CTkFrame(tab)
        inp.grid(row=2, column=0, padx=8, pady=8, sticky="ew")
        inp.grid_columnconfigure(0, weight=1)
        
        inp_inner = ctk.CTkFrame(inp, fg_color="transparent")
        inp_inner.grid(row=0, column=0, padx=(10, 4), pady=10, sticky="ew")
        inp_inner.grid_columnconfigure(0, weight=1)
        
        self.query_entry = ctk.CTkEntry(
            inp_inner, placeholder_text="검색어 입력 (Enter)",
            font=("", 13), height=38, state="disabled"
        )
        self.query_entry.grid(row=0, column=0, sticky="ew")
        self.query_entry.bind("<Return>", lambda e: self._search())
        
        ctk.CTkButton(
            inp_inner, text="🕐", width=34, height=34,
            command=self._show_history
        ).grid(row=0, column=1, padx=(4, 0))
        
        ctk.CTkButton(
            inp, text="✕", width=34, height=34,
            command=lambda: self.query_entry.delete(0, 'end')
        ).grid(row=0, column=1, padx=3, pady=10)
        
        ctk.CTkLabel(inp, text="결과:").grid(row=0, column=2, padx=(10, 4))
        
        self.k_entry = ctk.CTkEntry(inp, width=45, height=34)
        self.k_entry.grid(row=0, column=3, padx=3, pady=10)
        self.k_entry.insert(0, str(AppConfig.DEFAULT_SEARCH_RESULTS))
        
        self.search_btn = ctk.CTkButton(
            inp, text="🔍 검색", width=90, height=38,
            command=self._search, state="disabled"
        )
        self.search_btn.grid(row=0, column=4, padx=(4, 10), pady=10)
    
    def _build_files_tab(self):
        tab = self.tabs.tab("📄 파일")
        self.file_panel = FileListPanel(tab)
        self.file_panel.pack(fill="both", expand=True, padx=8, pady=8)
    
    def _build_settings_tab(self):
        tab = self.tabs.tab("⚙️ 설정")
        
        # 테마
        theme_f = ctk.CTkFrame(tab)
        theme_f.pack(padx=15, pady=12, fill="x")
        ctk.CTkLabel(theme_f, text="🎨 테마", font=ctk.CTkFont(size=13, weight="bold")).pack(pady=(10, 8), anchor="w", padx=12)
        ctk.CTkOptionMenu(theme_f, values=["System", "Light", "Dark"], command=ctk.set_appearance_mode, width=180).pack(pady=(0, 10), anchor="w", padx=12)
        
        # 검색
        search_f = ctk.CTkFrame(tab)
        search_f.pack(padx=15, pady=12, fill="x")
        ctk.CTkLabel(search_f, text="🔍 검색", font=ctk.CTkFont(size=13, weight="bold")).pack(pady=(10, 8), anchor="w", padx=12)
        
        self.hybrid_sw = ctk.CTkSwitch(search_f, text="하이브리드 검색", command=self._toggle_hybrid)
        self.hybrid_sw.pack(pady=(0, 4), anchor="w", padx=12)
        if self.hybrid:
            self.hybrid_sw.select()
        
        ctk.CTkLabel(search_f, text="벡터 + 키워드 결합 검색", text_color="gray", font=("", 10)).pack(pady=(0, 10), anchor="w", padx=12)
        
        # 모델
        model_f = ctk.CTkFrame(tab)
        model_f.pack(padx=15, pady=12, fill="x")
        ctk.CTkLabel(model_f, text="🤖 AI 모델", font=ctk.CTkFont(size=13, weight="bold")).pack(pady=(10, 8), anchor="w", padx=12)
        
        self.model_menu = ctk.CTkOptionMenu(
            model_f, values=list(AppConfig.AVAILABLE_MODELS.keys()),
            command=self._on_model, width=220
        )
        self.model_menu.set(self.model_name)
        self.model_menu.pack(pady=(0, 4), anchor="w", padx=12)
        
        ctk.CTkLabel(model_f, text="⚠️ 변경 시 재시작 필요", text_color="orange", font=("", 10)).pack(pady=(0, 10), anchor="w", padx=12)
        
        # 데이터
        data_f = ctk.CTkFrame(tab)
        data_f.pack(padx=15, pady=12, fill="x")
        ctk.CTkLabel(data_f, text="🗂️ 데이터", font=ctk.CTkFont(size=13, weight="bold")).pack(pady=(10, 8), anchor="w", padx=12)
        
        btn_row = ctk.CTkFrame(data_f, fg_color="transparent")
        btn_row.pack(anchor="w", padx=12, pady=(0, 10))
        
        ctk.CTkButton(btn_row, text="🗑️ 캐시삭제", command=self._clear_cache, fg_color="#DC2626", hover_color="#B91C1C", width=100).pack(side="left", padx=(0, 8))
        ctk.CTkButton(btn_row, text="🕐 히스토리삭제", command=self._clear_history, width=115).pack(side="left")
        
        # 키보드 단축키
        key_f = ctk.CTkFrame(tab)
        key_f.pack(padx=15, pady=12, fill="x")
        ctk.CTkLabel(key_f, text="⌨️ 키보드 단축키", font=ctk.CTkFont(size=13, weight="bold")).pack(pady=(10, 8), anchor="w", padx=12)
        
        shortcuts = [
            ("Ctrl+O", "폴더 열기"),
            ("Ctrl+F", "검색창 포커스"),
            ("Ctrl+H", "검색 기록"),
            ("Ctrl+S", "기록 저장"),
            ("Esc", "검색어 지우기"),
        ]
        
        for key, desc in shortcuts:
            row = ctk.CTkFrame(key_f, fg_color="transparent")
            row.pack(fill="x", padx=12, pady=1)
            ctk.CTkLabel(row, text=key, font=ctk.CTkFont(size=11, weight="bold"), width=80, anchor="w").pack(side="left")
            ctk.CTkLabel(row, text=desc, font=ctk.CTkFont(size=11), text_color="gray").pack(side="left")
    
    def _build_statusbar(self):
        bar = ctk.CTkFrame(self, height=30)
        bar.grid(row=1, column=0, padx=12, pady=(0, 8), sticky="ew")
        
        self.status_lbl = ctk.CTkLabel(bar, text="🔄 초기화 중...", anchor="w")
        self.status_lbl.pack(side="left", padx=12, pady=6)
        
        ctk.CTkLabel(bar, text=f"v{AppConfig.APP_VERSION}", text_color="gray").pack(side="right", padx=12, pady=6)
    
    # ========================================================================
    # 초기화
    # ========================================================================
    def _init_app(self):
        self.status = TaskStatus.LOADING_MODEL
        self._set_status("모델 로딩 중...", "orange")
        
        def cb(msg):
            self.after(0, lambda: self._set_status(msg, "orange"))
        
        def task():
            res = self.qa.load_model(self.model_name, cb)
            self.after(0, lambda: self._on_model_loaded(res))
        
        threading.Thread(target=task, daemon=True).start()
    
    def _on_model_loaded(self, res: TaskResult):
        self.status = TaskStatus.IDLE
        
        if not res.success:
            messagebox.showerror("오류", res.message)
            self._set_status("모델 로드 실패", "red")
            return
        
        self._set_status(res.message, "green")
        self.folder_btn.configure(state="normal")
        
        if self.last_folder and os.path.isdir(self.last_folder):
            self.recent_btn.configure(state="normal")
        
        self._toast("준비 완료", "success")
        self._show_welcome()
    
    def _show_welcome(self):
        for w in self.result_scroll.winfo_children():
            w.destroy()
        
        ctk.CTkLabel(
            self.result_scroll,
            text="👋 폴더를 선택하고 검색을 시작하세요",
            text_color="gray", font=ctk.CTkFont(size=13)
        ).pack(pady=80)
    
    # ========================================================================
    # 폴더
    # ========================================================================
    def _open_folder(self):
        folder = filedialog.askdirectory(title="규정 폴더 선택")
        if folder:
            self._load_folder(folder)
    
    def _load_recent(self):
        if self.last_folder and os.path.isdir(self.last_folder):
            self._load_folder(self.last_folder)
        else:
            messagebox.showwarning("경고", "최근 폴더 없음")
    
    def _refresh(self):
        if self.last_folder:
            cache = self.qa._get_cache_dir(self.last_folder)
            shutil.rmtree(cache, ignore_errors=True)
            self._load_folder(self.last_folder)
    
    def _load_folder(self, folder: str):
        if self.status != TaskStatus.IDLE:
            messagebox.showwarning("경고", "처리 중입니다")
            return
        
        try:
            files = [
                os.path.join(folder, f)
                for f in os.listdir(folder)
                if os.path.splitext(f)[1].lower() in AppConfig.SUPPORTED_EXTENSIONS
            ]
        except Exception as e:
            messagebox.showerror("오류", f"폴더 접근 실패: {e}")
            return
        
        if not files:
            messagebox.showwarning("경고", "지원 파일 없음")
            return
        
        self.status = TaskStatus.PROCESSING_DOCS
        self.folder_lbl.configure(text=folder, text_color=("#1E3A5F", "#E0E0E0"))
        self._set_ui(False)
        
        dlg = ProgressDlg(self, "문서 처리")
        
        def cb(pct, msg):
            self.after(0, lambda: dlg.update(pct, msg))
        
        def task():
            res = self.qa.process_documents(folder, files, cb)
            self.after(0, lambda: self._on_folder_done(res, folder, dlg))
        
        threading.Thread(target=task, daemon=True).start()
    
    def _on_folder_done(self, res: TaskResult, folder: str, dlg: ProgressDlg):
        dlg.destroy()
        self.status = TaskStatus.IDLE
        self._set_ui(True)
        
        self.file_panel.update_files(self.qa.get_file_infos())
        
        if res.success:
            self.last_folder = folder
            self._save_config()
            
            self._set_status(res.message, "green")
            self.query_entry.configure(state="normal")
            self.search_btn.configure(state="normal")
            self.refresh_btn.configure(state="normal")
            self.recent_btn.configure(state="normal")
            
            for w in self.result_scroll.winfo_children():
                w.destroy()
            
            info = f"✅ 준비 완료\n\n"
            if res.data:
                info += f"청크: {res.data.get('chunks', 0)} / 신규: {res.data.get('new', 0)} / 캐시: {res.data.get('cached', 0)}"
            
            ctk.CTkLabel(
                self.result_scroll, text=info,
                font=ctk.CTkFont(size=12), justify="center"
            ).pack(pady=60)
            
            self.query_entry.focus()
            
            if res.failed_items:
                self._show_failed(res.failed_items)
            else:
                self._toast("문서 처리 완료", "success")
        else:
            self._set_status(res.message, "red")
            messagebox.showerror("오류", res.message)
    
    def _show_failed(self, items: List[str]):
        msg = '\n'.join(f"• {f}" for f in items[:8])
        if len(items) > 8:
            msg += f"\n... 외 {len(items)-8}개"
        messagebox.showwarning("일부 실패", msg)
    
    # ========================================================================
    # 검색
    # ========================================================================
    def _search(self):
        query = self.query_entry.get().strip()
        
        if not query:
            self._toast("검색어 입력", "warning")
            return
        
        if self.status != TaskStatus.IDLE:
            return
        
        if not self.qa.vector_store:
            self._toast("문서 로드 필요", "warning")
            return
        
        try:
            k = int(self.k_entry.get())
            k = max(1, min(k, AppConfig.MAX_SEARCH_RESULTS))
        except:
            k = AppConfig.DEFAULT_SEARCH_RESULTS
        
        self.status = TaskStatus.SEARCHING
        self._set_search_ui(False)
        
        for w in self.result_scroll.winfo_children():
            w.destroy()
        ctk.CTkLabel(self.result_scroll, text="🔍 검색 중...", font=ctk.CTkFont(size=13)).pack(pady=80)
        
        def task():
            res = self.qa.search(query, k, self.hybrid)
            self.after(0, lambda: self._on_search_done(res, query))
        
        threading.Thread(target=task, daemon=True).start()
    
    def _on_search_done(self, res: TaskResult, query: str):
        self.status = TaskStatus.IDLE
        self._set_search_ui(True)
        
        for w in self.result_scroll.winfo_children():
            w.destroy()
        
        if not res.success:
            ctk.CTkLabel(self.result_scroll, text=f"❌ {res.message}", text_color="red").pack(pady=80)
            return
        
        if not res.data:
            ctk.CTkLabel(self.result_scroll, text="❌ 결과 없음", text_color="gray").pack(pady=80)
            return
        
        self.history.add(query)
        
        # 쿼리 표시
        q_frame = ctk.CTkFrame(self.result_scroll)
        q_frame.pack(fill="x", padx=4, pady=(4, 12))
        
        ctk.CTkLabel(
            q_frame, text=f"🔎 {query}",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(pady=8, padx=12, anchor="w")
        
        mode = "하이브리드" if self.hybrid else "벡터"
        ctk.CTkLabel(
            q_frame, text=f"{mode} | {len(res.data)}개 결과",
            text_color="gray", font=ctk.CTkFont(size=10)
        ).pack(pady=(0, 8), padx=12, anchor="w")
        
        # 결과
        for i, item in enumerate(res.data, 1):
            ResultCard(
                self.result_scroll, i, item,
                on_copy=self._copy,
                font_size=self.font_size,
                query=query
            ).pack(fill="x", padx=4, pady=4)
        
        self.chat_log.append({'role': 'user', 'content': query})
        self.chat_log.append({'role': 'bot', 'content': f"{len(res.data)}개 결과"})
        
        self.query_entry.delete(0, 'end')
        self.query_entry.focus()
    
    def _copy(self, text: str):
        self.clipboard_clear()
        self.clipboard_append(text)
        self._toast("복사됨", "success")
    
    def _show_history(self):
        items = self.history.get(8)
        if not items:
            self._toast("검색 기록 없음", "info")
            return
        
        popup = HistoryPopup(self, items, self._select_history)
        popup.update_idletasks()
        x = self.query_entry.winfo_rootx()
        y = self.query_entry.winfo_rooty() - popup.winfo_height() - 5
        popup.geometry(f"+{x}+{y}")
        popup.focus()
    
    def _select_history(self, q: str):
        self.query_entry.delete(0, 'end')
        self.query_entry.insert(0, q)
        self.query_entry.focus()
    
    # ========================================================================
    # 헬퍼
    # ========================================================================
    def _set_status(self, text: str, color: str = "white"):
        self.status_lbl.configure(text=text, text_color=color)
    
    def _set_ui(self, enabled: bool):
        state = "normal" if enabled else "disabled"
        self.folder_btn.configure(state=state)
        if self.last_folder:
            self.recent_btn.configure(state=state)
    
    def _set_search_ui(self, enabled: bool):
        state = "normal" if enabled else "disabled"
        self.query_entry.configure(state=state)
        self.search_btn.configure(state=state)
    
    def _toast(self, msg: str, type_: str = "info"):
        Toast(self, msg, type_=type_)
    
    def _font(self, delta: int):
        new = self.font_size + delta
        if AppConfig.MIN_FONT_SIZE <= new <= AppConfig.MAX_FONT_SIZE:
            self.font_size = new
            self.font_lbl.configure(text=str(new))
    
    def _toggle_hybrid(self):
        self.hybrid = self.hybrid_sw.get() == 1
        self._save_config()
    
    def _on_model(self, name: str):
        self.model_name = name
        self._save_config()
    
    def _reset(self):
        self.chat_log.clear()
        self._toast("초기화됨", "info")
    
    def _save_log(self):
        if not self.chat_log:
            messagebox.showwarning("경고", "저장할 기록 없음")
            return
        
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text", "*.txt"), ("JSON", "*.json")],
            initialfile=f"검색기록_{datetime.now():%Y%m%d_%H%M%S}"
        )
        
        if not path:
            return
        
        try:
            if path.endswith('.json'):
                # JSON 형식 저장
                export_data = {
                    "app": AppConfig.APP_NAME,
                    "version": AppConfig.APP_VERSION,
                    "exported_at": datetime.now().isoformat(),
                    "folder": self.last_folder,
                    "logs": self.chat_log
                }
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(export_data, f, ensure_ascii=False, indent=2)
            else:
                # 텍스트 형식 저장
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(f"{'='*60}\n{AppConfig.APP_NAME} 검색 기록\n{'='*60}\n\n")
                    for item in self.chat_log:
                        f.write(f"{'👤' if item['role']=='user' else '🔍'} {item['content']}\n{'-'*60}\n\n")
            self._toast("저장됨", "success")
        except Exception as e:
            messagebox.showerror("오류", str(e))
    
    def _clear_cache(self):
        if messagebox.askyesno("확인", "캐시를 삭제하시겠습니까?"):
            res = self.qa.clear_cache()
            self._toast(res.message, "success" if res.success else "error")
    
    def _clear_history(self):
        if messagebox.askyesno("확인", "히스토리를 삭제하시겠습니까?"):
            self.history.clear()
            self._toast("삭제됨", "success")
    
    def _on_close(self):
        self._save_config()
        self.qa.cleanup()
        logger.info("종료")
        self.destroy()


# ============================================================================
# 메인
# ============================================================================
def main():
    try:
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")
        
        logger.info(f"시작 v{AppConfig.APP_VERSION}")
        
        qa = RegulationQASystem()
        app = App(qa)
        app.mainloop()
        
    except Exception as e:
        logger.critical(f"치명적 오류: {e}", exc_info=True)
        try:
            messagebox.showerror("오류", str(e))
        except:
            pass
        sys.exit(1)


if __name__ == "__main__":
    main()
