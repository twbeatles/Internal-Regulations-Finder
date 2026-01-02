# -*- coding: utf-8 -*-
"""
사내 규정 검색기 v9.0 (PyQt6)
깔끔하고 전문적인 UI
"""

from __future__ import annotations
import sys
import os
import json
import threading
import tempfile
import hashlib
import shutil
import logging
import subprocess
import platform
import re
import gc
import math
from typing import List, Dict, Tuple, Optional, Callable, Any, TYPE_CHECKING
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum, auto
from collections import Counter

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QTabWidget, QFrame,
    QScrollArea, QProgressBar, QFileDialog, QMessageBox, QComboBox,
    QCheckBox, QSpinBox, QTableWidget, QTableWidgetItem, QHeaderView, QSlider,
    QSplitter, QStackedWidget, QSizePolicy, QGraphicsDropShadowEffect, QMenu
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QPropertyAnimation, QEasingCurve, QSize
from PyQt6.QtGui import QFont, QIcon, QColor, QPalette, QAction, QTextCursor, QTextCharFormat, QShortcut, QKeySequence

if TYPE_CHECKING:
    from langchain_community.vectorstores import FAISS
    from langchain_huggingface import HuggingFaceEmbeddings

# ============================================================================
# 상수 및 설정
# ============================================================================
class AppConfig:
    APP_NAME = "사내 규정 검색기"
    APP_VERSION = "9.0"
    
    AVAILABLE_MODELS: Dict[str, str] = {
        "SNU SBERT (고성능)": "snunlp/KR-SBERT-V40K-klueNLI-augSTS",
        "BM-K Simal (균형)": "BM-K/ko-simal-roberta-base",
        "JHGan SBERT (빠름)": "jhgan/ko-sbert-nli"
    }
    DEFAULT_MODEL = "JHGan SBERT (빠름)"
    
    CONFIG_FILE = "config.json"
    HISTORY_FILE = "search_history.json"
    SUPPORTED_EXTENSIONS = ('.txt', '.docx', '.pdf')
    
    MAX_FONT_SIZE = 32
    MIN_FONT_SIZE = 8
    DEFAULT_FONT_SIZE = 14
    DEFAULT_SEARCH_RESULTS = 3
    MAX_SEARCH_RESULTS = 10
    MAX_HISTORY_SIZE = 30
    
    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 80
    VECTOR_WEIGHT = 0.7
    BM25_WEIGHT = 0.3
    MAX_DOCS_IN_MEMORY = 5000
    BATCH_SIZE = 100


class TaskStatus(Enum):
    IDLE = auto()
    LOADING_MODEL = auto()
    PROCESSING_DOCS = auto()
    SEARCHING = auto()


class FileStatus(Enum):
    PENDING = "대기"
    PROCESSING = "처리중"
    SUCCESS = "완료"
    FAILED = "실패"
    CACHED = "캐시"


@dataclass
class TaskResult:
    success: bool
    message: str
    data: Any = None
    failed_items: List[str] = field(default_factory=list)


@dataclass
class FileInfo:
    path: str
    name: str
    extension: str
    size: int
    status: FileStatus = FileStatus.PENDING
    chunks: int = 0
    error: str = ""


# ============================================================================
# 유틸리티
# ============================================================================
def get_app_directory() -> str:
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def setup_logger() -> logging.Logger:
    logger = logging.getLogger('RegSearch')
    if logger.handlers:
        return logger
    logger.setLevel(logging.INFO)
    log_dir = os.path.join(get_app_directory(), 'logs')
    os.makedirs(log_dir, exist_ok=True)
    fh = logging.FileHandler(
        os.path.join(log_dir, f'app_{datetime.now():%Y%m%d}.log'),
        encoding='utf-8'
    )
    fh.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(fh)
    if not getattr(sys, 'frozen', False):
        ch = logging.StreamHandler()
        ch.setLevel(logging.INFO)
        logger.addHandler(ch)
    return logger

logger = setup_logger()


class FileUtils:
    @staticmethod
    def safe_read(path: str, encoding: str = 'utf-8') -> Tuple[Optional[str], Optional[str]]:
        try:
            with open(path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read(), None
        except Exception as e:
            return None, str(e)
    
    @staticmethod
    def get_metadata(path: str) -> Optional[Dict]:
        try:
            stat = os.stat(path)
            return {'size': stat.st_size, 'mtime': stat.st_mtime}
        except OSError as e:
            logger.debug(f"파일 메타데이터 조회 실패: {path} - {e}")
            return None
    
    @staticmethod
    def open_file(path: str):
        try:
            if platform.system() == 'Windows':
                os.startfile(path)
            elif platform.system() == 'Darwin':
                subprocess.run(['open', path], check=False)
            else:
                subprocess.run(['xdg-open', path], check=False)
        except Exception as e:
            logger.error(f"파일 열기 실패: {e}")
    
    @staticmethod
    def format_size(size: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f}{unit}"
            size /= 1024
        return f"{size:.1f}TB"


# ============================================================================
# BM25
# ============================================================================
class BM25Light:
    __slots__ = ['k1', 'b', 'corpus', 'doc_lens', 'avgdl', 'idf', 'N']
    
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus: List[List[str]] = []
        self.doc_lens: List[int] = []
        self.avgdl = 0.0
        self.idf: Dict[str, float] = {}
        self.N = 0
    
    def _tokenize(self, text: str) -> List[str]:
        if not text:
            return []
        text = re.sub(r'[^\w\s가-힣]', ' ', text.lower())
        return [t for t in text.split() if len(t) >= 2]
    
    def fit(self, docs: List[str]):
        self.corpus = []
        self.doc_lens = []
        df = Counter()
        for doc in docs:
            tokens = self._tokenize(doc)
            self.corpus.append(tokens)
            self.doc_lens.append(len(tokens))
            df.update(set(tokens))
        self.N = len(docs)
        self.avgdl = sum(self.doc_lens) / self.N if self.N else 0
        self.idf = {t: math.log((self.N - f + 0.5) / (f + 0.5) + 1) for t, f in df.items()}
        del df
        gc.collect()
    
    def search(self, query: str, top_k: int = 5) -> List[Tuple[int, float]]:
        if not self.corpus or not query:
            return []
        q_tokens = self._tokenize(query)
        if not q_tokens:
            return []
        scores = []
        for idx, doc_tokens in enumerate(self.corpus):
            if not doc_tokens:
                continue
            score = self._score(q_tokens, doc_tokens, self.doc_lens[idx])
            if score > 0:
                scores.append((idx, score))
        scores.sort(key=lambda x: x[1], reverse=True)
        return scores[:top_k]
    
    def _score(self, query: List[str], doc: List[str], doc_len: int) -> float:
        score = 0.0
        doc_tf = Counter(doc)
        for term in query:
            if term not in self.idf:
                continue
            tf = doc_tf.get(term, 0)
            idf = self.idf[term]
            num = tf * (self.k1 + 1)
            den = tf + self.k1 * (1 - self.b + self.b * doc_len / self.avgdl)
            score += idf * num / den if den > 0 else 0
        return score
    
    def clear(self):
        self.corpus.clear()
        self.doc_lens.clear()
        self.idf.clear()
        gc.collect()


# ============================================================================
# 문서 추출기
# ============================================================================
class DocumentExtractor:
    def __init__(self):
        self._docx_module = None
        self._pdf_module = None
    
    @property
    def docx(self):
        if self._docx_module is None:
            try:
                from docx import Document
                self._docx_module = Document
            except ImportError:
                self._docx_module = False
        return self._docx_module
    
    @property
    def pdf(self):
        if self._pdf_module is None:
            try:
                from pypdf import PdfReader
                self._pdf_module = PdfReader
            except ImportError:
                self._pdf_module = False
        return self._pdf_module
    
    def extract(self, path: str) -> Tuple[str, Optional[str]]:
        if not path or not os.path.exists(path):
            return "", f"파일 없음: {path}"
        if not os.path.isfile(path):
            return "", f"파일이 아님: {path}"
        ext = os.path.splitext(path)[1].lower()
        if ext == '.txt':
            return self._extract_txt(path)
        elif ext == '.docx':
            return self._extract_docx(path)
        elif ext == '.pdf':
            return self._extract_pdf(path)
        return "", f"지원하지 않는 형식: {ext}"
    
    def _extract_txt(self, path: str) -> Tuple[str, Optional[str]]:
        return FileUtils.safe_read(path)
    
    def _extract_docx(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.docx:
            return "", "DOCX 라이브러리 없음"
        try:
            doc = self.docx(path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            return '\n\n'.join(parts), None
        except Exception as e:
            return "", f"DOCX 오류: {e}"
    
    def _extract_pdf(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.pdf:
            return "", "PDF 라이브러리 없음"
        try:
            reader = self.pdf(path)
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception as e:
                    return "", "암호화된 PDF"
            texts = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
                except Exception:
                    continue
            if not texts:
                return "", "텍스트 없음 (이미지 PDF)"
            return '\n\n'.join(texts), None
        except Exception as e:
            return "", f"PDF 오류: {e}"


# ============================================================================
# QA 시스템
# ============================================================================
class RegulationQASystem:
    def __init__(self):
        self.vector_store = None
        self.embedding_model = None
        self.model_id = None
        self.extractor = DocumentExtractor()
        self.cache_path = os.path.join(tempfile.gettempdir(), "reg_qa_v90")
        self.bm25 = None
        self.documents: List[str] = []
        self.doc_meta: List[Dict] = []
        self.file_infos: Dict[str, FileInfo] = {}
        self.current_folder = ""
        self._lock = threading.Lock()
    
    def load_model(self, model_name: str, progress_cb=None) -> TaskResult:
        model_id = AppConfig.AVAILABLE_MODELS.get(model_name, AppConfig.AVAILABLE_MODELS[AppConfig.DEFAULT_MODEL])
        try:
            if progress_cb: progress_cb("라이브러리 로드 중...")
            import torch
            from langchain_huggingface import HuggingFaceEmbeddings
            if progress_cb: progress_cb("모델 로딩 중...")
            cache_dir = os.path.join(get_app_directory(), 'models')
            os.makedirs(cache_dir, exist_ok=True)
            device = 'cuda' if torch.cuda.is_available() else 'cpu'
            self.embedding_model = HuggingFaceEmbeddings(
                model_name=model_id, cache_folder=cache_dir,
                model_kwargs={'device': device}, encode_kwargs={'normalize_embeddings': True}
            )
            self.model_id = model_id
            gc.collect()
            if device == 'cuda': torch.cuda.empty_cache()
            logger.info(f"모델 로드 완료: {model_name} ({device})")
            return TaskResult(True, f"모델 로드 완료 ({device})")
        except Exception as e:
            logger.error(f"모델 로드 실패: {e}")
            return TaskResult(False, f"모델 로드 실패: {e}")
    
    def _get_cache_dir(self, folder: str) -> str:
        if not self.model_id:
            raise ValueError("모델이 로드되지 않았습니다")
        h1 = hashlib.md5(self.model_id.encode()).hexdigest()[:6]
        h2 = hashlib.md5(folder.encode()).hexdigest()[:6]
        return os.path.join(self.cache_path, f"{h2}_{h1}")
    
    def process_documents(self, folder: str, files: List[str], progress_cb) -> TaskResult:
        if not self.embedding_model:
            return TaskResult(False, "모델이 로드되지 않았습니다")
        with self._lock:
            return self._process_internal(folder, files, progress_cb)
    
    def _process_internal(self, folder, files, progress_cb) -> TaskResult:
        from langchain.text_splitter import CharacterTextSplitter
        from langchain_community.vectorstores import FAISS
        from langchain.docstore.document import Document
        
        self.current_folder = folder
        cache_dir = self._get_cache_dir(folder)
        self.file_infos.clear()
        
        for fp in files:
            meta = FileUtils.get_metadata(fp)
            self.file_infos[fp] = FileInfo(fp, os.path.basename(fp), os.path.splitext(fp)[1].lower(), meta['size'] if meta else 0)
        
        progress_cb(5, "캐시 확인...")
        cache_info = self._load_cache_info(cache_dir)
        to_process, cached = [], []
        
        for fp in files:
            fname = os.path.basename(fp)
            meta = FileUtils.get_metadata(fp)
            if meta and fname in cache_info:
                cm = cache_info[fname]
                if cm.get('size') == meta['size'] and cm.get('mtime') == meta['mtime']:
                    cached.append(fp)
                    self.file_infos[fp].status = FileStatus.CACHED
                    self.file_infos[fp].chunks = cm.get('chunks', 0)
                    continue
            to_process.append(fp)
        
        self.documents, self.doc_meta = [], []
        
        if cached and os.path.exists(os.path.join(cache_dir, "index.faiss")):
            try:
                progress_cb(10, "캐시 로드...")
                self.vector_store = FAISS.load_local(cache_dir, self.embedding_model, allow_dangerous_deserialization=True)
                docs_path = os.path.join(cache_dir, "docs.json")
                if os.path.exists(docs_path):
                    with open(docs_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        self.documents = data.get('docs', [])
                        self.doc_meta = data.get('meta', [])
            except Exception as e:
                logger.warning(f"캐시 로드 실패: {e}")
                to_process, cached = files, []
                self.vector_store = None
        
        if not to_process:
            self._build_bm25()
            progress_cb(100, "완료!")
            return TaskResult(True, f"캐시에서 {len(cached)}개 파일 로드", {'chunks': len(self.documents), 'cached': len(cached), 'new': 0})
        
        splitter = CharacterTextSplitter(separator="\n\n", chunk_size=AppConfig.CHUNK_SIZE, chunk_overlap=AppConfig.CHUNK_OVERLAP)
        failed, new_docs, new_cache_info = [], [], {}
        
        for i, fp in enumerate(to_process):
            fname = os.path.basename(fp)
            progress_cb(15 + int((i / len(to_process)) * 55), f"처리: {fname}")
            self.file_infos[fp].status = FileStatus.PROCESSING
            try:
                content, error = self.extractor.extract(fp)
                if error:
                    failed.append(f"{fname} ({error})")
                    self.file_infos[fp].status = FileStatus.FAILED
                    continue
                if not content.strip():
                    failed.append(f"{fname} (빈 파일)")
                    self.file_infos[fp].status = FileStatus.FAILED
                    continue
                chunks = splitter.split_text(content)
                chunk_count = 0
                for chunk in chunks:
                    if chunk.strip():
                        new_docs.append(Document(page_content=chunk.strip(), metadata={"source": fname, "path": fp}))
                        self.documents.append(chunk.strip())
                        self.doc_meta.append({"source": fname, "path": fp})
                        chunk_count += 1
                self.file_infos[fp].status = FileStatus.SUCCESS
                self.file_infos[fp].chunks = chunk_count
                meta = FileUtils.get_metadata(fp)
                if meta:
                    new_cache_info[fname] = {'size': meta['size'], 'mtime': meta['mtime'], 'chunks': chunk_count}
            except Exception as e:
                failed.append(f"{fname} ({e})")
                self.file_infos[fp].status = FileStatus.FAILED
        
        if not new_docs and not self.vector_store:
            return TaskResult(False, "처리 가능한 문서 없음", failed_items=failed)
        
        progress_cb(75, "벡터 인덱스 생성...")
        try:
            if new_docs:
                if self.vector_store:
                    for i in range(0, len(new_docs), AppConfig.BATCH_SIZE):
                        self.vector_store.add_documents(new_docs[i:i + AppConfig.BATCH_SIZE])
                else:
                    self.vector_store = FAISS.from_documents(new_docs, self.embedding_model)
        except Exception as e:
            return TaskResult(False, f"인덱스 생성 실패: {e}")
        
        progress_cb(85, "키워드 인덱스 생성...")
        self._build_bm25()
        progress_cb(90, "캐시 저장...")
        self._save_cache(cache_dir, cache_info, new_cache_info)
        gc.collect()
        progress_cb(100, "완료!")
        return TaskResult(True, f"{len(files) - len(failed)}개 처리 완료", {'chunks': len(self.documents), 'new': len(to_process) - len(failed), 'cached': len(cached)}, failed)
    
    def _build_bm25(self):
        if self.documents:
            self.bm25 = BM25Light()
            self.bm25.fit(self.documents)
    
    def _load_cache_info(self, cache_dir):
        path = os.path.join(cache_dir, "cache_info.json")
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except: pass
        return {}
    
    def _save_cache(self, cache_dir, old_info, new_info):
        try:
            os.makedirs(cache_dir, exist_ok=True)
            self.vector_store.save_local(cache_dir)
            with open(os.path.join(cache_dir, "cache_info.json"), 'w', encoding='utf-8') as f:
                json.dump({**old_info, **new_info}, f, ensure_ascii=False)
            with open(os.path.join(cache_dir, "docs.json"), 'w', encoding='utf-8') as f:
                json.dump({'docs': self.documents, 'meta': self.doc_meta}, f, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"캐시 저장 실패: {e}")
    
    def search(self, query: str, k: int = 3, hybrid: bool = True) -> TaskResult:
        if not self.vector_store:
            return TaskResult(False, "문서가 로드되지 않음")
        query = query.strip()
        if len(query) < 2:
            return TaskResult(False, "검색어가 너무 짧습니다")
        try:
            k = max(1, min(k, AppConfig.MAX_SEARCH_RESULTS))
            vec_results = self.vector_store.similarity_search_with_score(query, k=k*2)
            results = {}
            if vec_results:
                min_d, max_d = min(r[1] for r in vec_results), max(r[1] for r in vec_results)
                rng = max_d - min_d if max_d != min_d else 1
                for doc, dist in vec_results:
                    key = doc.page_content[:100]
                    score = max(0.1, 1 - ((dist - min_d) / (rng + 0.001)))
                    results[key] = {'content': doc.page_content, 'source': doc.metadata.get('source', '?'), 'path': doc.metadata.get('path', ''), 'vec_score': score, 'bm25_score': 0}
            if hybrid and self.bm25:
                bm_res = self.bm25.search(query, top_k=k*2)
                if bm_res:
                    max_bm = max(r[1] for r in bm_res) if bm_res else 1
                    for idx, sc in bm_res:
                        if 0 <= idx < len(self.documents):
                            key = self.documents[idx][:100]
                            norm = sc / (max_bm + 0.001)
                            if key in results:
                                results[key]['bm25_score'] = norm
                            else:
                                meta = self.doc_meta[idx] if idx < len(self.doc_meta) else {}
                                results[key] = {'content': self.documents[idx], 'source': meta.get('source', '?'), 'path': meta.get('path', ''), 'vec_score': 0, 'bm25_score': norm}
            for item in results.values():
                item['score'] = AppConfig.VECTOR_WEIGHT * item['vec_score'] + AppConfig.BM25_WEIGHT * item['bm25_score']
            sorted_res = sorted(results.values(), key=lambda x: x['score'], reverse=True)[:k]
            return TaskResult(True, "검색 완료", sorted_res)
        except Exception as e:
            logger.error(f"검색 오류: {e}")
            return TaskResult(False, f"검색 오류: {e}")
    
    def get_file_infos(self): return list(self.file_infos.values())
    def clear_cache(self):
        if os.path.exists(self.cache_path):
            shutil.rmtree(self.cache_path, ignore_errors=True)
        return TaskResult(True, "캐시 삭제 완료")
    def cleanup(self):
        self.documents.clear()
        self.doc_meta.clear()
        if self.bm25: self.bm25.clear()
        gc.collect()


# ============================================================================
# 워커 스레드
# ============================================================================
class ModelLoaderThread(QThread):
    progress = pyqtSignal(str)
    finished = pyqtSignal(object)
    
    def __init__(self, qa, model_name):
        super().__init__()
        self.qa = qa
        self.model_name = model_name
    
    def run(self):
        result = self.qa.load_model(self.model_name, lambda msg: self.progress.emit(msg))
        self.finished.emit(result)


class DocumentProcessorThread(QThread):
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(object)
    
    def __init__(self, qa, folder, files):
        super().__init__()
        self.qa = qa
        self.folder = folder
        self.files = files
    
    def run(self):
        result = self.qa.process_documents(self.folder, self.files, lambda p, m: self.progress.emit(p, m))
        self.finished.emit(result)


class SearchThread(QThread):
    finished = pyqtSignal(object)
    
    def __init__(self, qa, query, k, hybrid):
        super().__init__()
        self.qa = qa
        self.query = query
        self.k = k
        self.hybrid = hybrid
    
    def run(self):
        result = self.qa.search(self.query, self.k, self.hybrid)
        self.finished.emit(result)


# ============================================================================
# 스타일
# ============================================================================
DARK_STYLE = """
/* 기본 위젯 */
QMainWindow, QWidget { background-color: #1a1a2e; color: #eaeaea; }
QLabel { color: #eaeaea; }

/* 탭 위젯 */
QTabWidget::pane { border: none; background: #16213e; border-radius: 8px; }
QTabBar::tab { background: #0f3460; color: #aaa; padding: 12px 24px; margin-right: 2px; border-top-left-radius: 8px; border-top-right-radius: 8px; font-weight: bold; }
QTabBar::tab:selected { background: #16213e; color: white; }
QTabBar::tab:hover:!selected { background: #1a4a70; color: #ddd; }

/* 입력 필드 */
QLineEdit { background: #0f3460; border: 2px solid #1a1a2e; border-radius: 8px; padding: 10px 15px; color: white; font-size: 14px; selection-background-color: #e94560; }
QLineEdit:focus { border-color: #e94560; }
QLineEdit:disabled { background: #2a2a3e; color: #666; }

/* 버튼 */
QPushButton { background: #0f3460; color: white; border: none; border-radius: 6px; padding: 10px 20px; font-weight: bold; }
QPushButton:hover { background: #e94560; }
QPushButton:pressed { background: #c73a52; }
QPushButton:disabled { background: #2a2a3e; color: #666; }

/* 텍스트 에디터 */
QTextEdit { background: #0f3460; border: none; border-radius: 6px; padding: 10px; color: #eaeaea; selection-background-color: #e94560; }

/* 프로그레스 바 */
QProgressBar { background: #0f3460; border: none; border-radius: 4px; height: 8px; text-align: center; }
QProgressBar::chunk { background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #e94560, stop:1 #0f3460); border-radius: 4px; }

/* 스크롤 영역 */
QScrollArea { border: none; background: transparent; }

/* 콤보박스 */
QComboBox { background: #0f3460; border: 2px solid #1a1a2e; border-radius: 6px; padding: 8px 12px; color: white; min-width: 100px; }
QComboBox:hover { border-color: #e94560; }
QComboBox:focus { border-color: #e94560; }
QComboBox::drop-down { border: none; width: 30px; subcontrol-origin: padding; subcontrol-position: center right; }
QComboBox::down-arrow { image: none; border-left: 5px solid transparent; border-right: 5px solid transparent; border-top: 6px solid #e94560; margin-right: 10px; }
QComboBox QAbstractItemView { background: #0f3460; border: 2px solid #e94560; border-radius: 4px; selection-background-color: #e94560; color: white; padding: 4px; }

/* 스핀박스 */
QSpinBox { background: #0f3460; border: 2px solid #1a1a2e; border-radius: 6px; padding: 6px 10px; color: white; min-width: 80px; }
QSpinBox:hover { border-color: #e94560; }
QSpinBox:focus { border-color: #e94560; }
QSpinBox::up-button, QSpinBox::down-button { background: #1a4a70; border: none; width: 20px; border-radius: 3px; }
QSpinBox::up-button:hover, QSpinBox::down-button:hover { background: #e94560; }
QSpinBox::up-arrow { image: none; border-left: 4px solid transparent; border-right: 4px solid transparent; border-bottom: 5px solid white; }
QSpinBox::down-arrow { image: none; border-left: 4px solid transparent; border-right: 4px solid transparent; border-top: 5px solid white; }

/* 체크박스 */
QCheckBox { color: #eaeaea; spacing: 8px; }
QCheckBox::indicator { width: 20px; height: 20px; border-radius: 4px; background: #0f3460; border: 2px solid #1a1a2e; }
QCheckBox::indicator:hover { border-color: #e94560; }
QCheckBox::indicator:checked { background: #e94560; border-color: #e94560; }

/* 테이블 */
QTableWidget { background: #16213e; border: none; gridline-color: #0f3460; alternate-background-color: #1a2845; }
QTableWidget::item { padding: 8px; border: none; }
QTableWidget::item:selected { background: #e94560; color: white; }
QTableWidget::item:hover:!selected { background: #1a4a70; }
QHeaderView::section { background: #0f3460; color: white; padding: 10px; border: none; font-weight: bold; }
QHeaderView::section:hover { background: #e94560; }

/* 스크롤바 - 세로 */
QScrollBar:vertical { background: #1a1a2e; width: 12px; border-radius: 6px; margin: 2px; }
QScrollBar::handle:vertical { background: #0f3460; border-radius: 5px; min-height: 30px; margin: 2px; }
QScrollBar::handle:vertical:hover { background: #e94560; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: none; }

/* 스크롤바 - 가로 */
QScrollBar:horizontal { background: #1a1a2e; height: 12px; border-radius: 6px; margin: 2px; }
QScrollBar::handle:horizontal { background: #0f3460; border-radius: 5px; min-width: 30px; margin: 2px; }
QScrollBar::handle:horizontal:hover { background: #e94560; }
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background: none; }

/* 툴팁 */
QToolTip { background: #16213e; color: white; border: 1px solid #e94560; padding: 8px 12px; border-radius: 6px; font-size: 12px; }

/* 메시지박스/다이얼로그 */
QMessageBox { background: #1a1a2e; }
QMessageBox QLabel { color: #eaeaea; }
QMessageBox QPushButton { min-width: 80px; }

/* 카드 프레임 */
QFrame#card { background: #16213e; border-radius: 12px; border: 1px solid #0f3460; }
QFrame#card:hover { border-color: #1a4a70; }
QFrame#statCard { background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #0f3460, stop:1 #16213e); border-radius: 12px; border: 1px solid #1a4a70; }
QFrame#resultCard { background: #16213e; border-radius: 12px; border: 1px solid #0f3460; }
QFrame#resultCard:hover { border-color: #e94560; }

/* 슬라이더 */
QSlider::groove:horizontal { background: #0f3460; height: 6px; border-radius: 3px; }
QSlider::handle:horizontal { background: #e94560; width: 16px; height: 16px; margin: -5px 0; border-radius: 8px; }
QSlider::handle:horizontal:hover { background: #ff6b8a; }
QSlider::sub-page:horizontal { background: #e94560; border-radius: 3px; }
"""


# ============================================================================
# 검색 히스토리
# ============================================================================
class SearchHistory:
    def __init__(self):
        self.items: List[str] = []
        self.path = os.path.join(get_app_directory(), AppConfig.HISTORY_FILE)
        self._load()
    
    def _load(self):
        if os.path.exists(self.path):
            try:
                with open(self.path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.items = [h['q'] for h in data[:AppConfig.MAX_HISTORY_SIZE]]
            except: self.items = []
    
    def _save(self):
        try:
            with open(self.path, 'w', encoding='utf-8') as f:
                json.dump([{'q': q} for q in self.items], f, ensure_ascii=False)
        except: pass
    
    def add(self, query: str):
        self.items = [q for q in self.items if q != query]
        self.items.insert(0, query)
        self.items = self.items[:AppConfig.MAX_HISTORY_SIZE]
        self._save()
    
    def get(self, count: int = 10): return self.items[:count]
    def clear(self): self.items = []; self._save()


# ============================================================================
# 결과 카드
# ============================================================================
class ResultCard(QFrame):
    """검색 결과를 표시하는 카드 위젯"""
    
    def __init__(self, idx: int, data: Dict, on_copy, font_size: int = 12, query: str = ""):
        super().__init__()
        self.setObjectName("resultCard")
        self.data = data
        self.query = query
        
        # 그림자 효과
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(15)
        shadow.setColor(QColor(0, 0, 0, 80))
        shadow.setOffset(0, 2)
        self.setGraphicsEffect(shadow)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 14, 18, 14)
        layout.setSpacing(10)
        
        # 헤더
        header = QHBoxLayout()
        header.setSpacing(12)
        
        # 결과 번호 배지
        badge = QLabel(f"{idx}")
        badge.setFixedSize(28, 28)
        badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        badge.setStyleSheet("""
            background: #e94560; 
            color: white; 
            border-radius: 14px; 
            font-weight: bold;
            font-size: 12px;
        """)
        header.addWidget(badge)
        
        # 파일명
        source = QLabel(data['source'])
        source.setStyleSheet("color: #e94560; font-size: 12px; font-weight: bold;")
        if data.get('path'):
            source.setToolTip(f"📁 {data['path']}\n더블클릭으로 파일 열기")
            source.setCursor(Qt.CursorShape.PointingHandCursor)
            source.mousePressEvent = lambda e: FileUtils.open_file(data['path']) if e.button() == Qt.MouseButton.LeftButton else None
        header.addWidget(source)
        
        header.addStretch()
        
        # 점수 표시
        score = int(data.get('score', 0) * 100)
        score_color = "#10b981" if score >= 70 else "#f59e0b" if score >= 40 else "#ef4444"
        
        score_container = QHBoxLayout()
        score_container.setSpacing(8)
        
        pbar = QProgressBar()
        pbar.setFixedWidth(80)
        pbar.setFixedHeight(8)
        pbar.setValue(score)
        pbar.setTextVisible(False)
        pbar.setStyleSheet(f"""
            QProgressBar {{ background: #0f3460; border-radius: 4px; }}
            QProgressBar::chunk {{ background: {score_color}; border-radius: 4px; }}
        """)
        score_container.addWidget(pbar)
        
        score_lbl = QLabel(f"{score}%")
        score_lbl.setStyleSheet(f"color: {score_color}; font-weight: bold; font-size: 13px;")
        score_lbl.setToolTip(f"유사도: {score}%\n벡터: {int(data.get('vec_score', 0)*100)}% | 키워드: {int(data.get('bm25_score', 0)*100)}%")
        score_container.addWidget(score_lbl)
        
        header.addLayout(score_container)
        
        # 버튼들
        btn_container = QHBoxLayout()
        btn_container.setSpacing(6)
        
        copy_btn = QPushButton("📋 복사")
        copy_btn.setFixedHeight(30)
        copy_btn.setFixedWidth(75)
        copy_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        copy_btn.clicked.connect(lambda: on_copy(data['content']))
        btn_container.addWidget(copy_btn)
        
        if data.get('path'):
            open_btn = QPushButton("📂 열기")
            open_btn.setFixedHeight(30)
            open_btn.setFixedWidth(75)
            open_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            open_btn.clicked.connect(lambda: FileUtils.open_file(data['path']))
            btn_container.addWidget(open_btn)
        
        header.addLayout(btn_container)
        layout.addLayout(header)
        
        # 내용 (검색어 하이라이트 포함)
        content = QTextEdit()
        content.setReadOnly(True)
        content.setFont(QFont("", font_size))
        content.setMinimumHeight(80)
        content.setMaximumHeight(180)
        
        # 검색어 하이라이트 적용
        self._apply_highlight(content, data['content'], query)
        
        layout.addWidget(content)
    
    def _apply_highlight(self, text_edit: QTextEdit, content: str, query: str):
        """검색어를 하이라이트 처리"""
        from PyQt6.QtGui import QTextCursor, QTextCharFormat
        
        text_edit.setPlainText(content)
        
        if not query or len(query) < 2:
            return
        
        # 검색어를 여러 단어로 분리하여 각각 하이라이트
        keywords = [k.strip() for k in query.split() if len(k.strip()) >= 2]
        
        highlight_format = QTextCharFormat()
        highlight_format.setBackground(QColor("#e94560"))
        highlight_format.setForeground(QColor("white"))
        
        cursor = text_edit.textCursor()
        
        for keyword in keywords:
            # 대소문자 무시 검색
            text = content.lower()
            keyword_lower = keyword.lower()
            start = 0
            
            while True:
                pos = text.find(keyword_lower, start)
                if pos == -1:
                    break
                
                cursor.setPosition(pos)
                cursor.setPosition(pos + len(keyword), QTextCursor.MoveMode.KeepAnchor)
                cursor.mergeCharFormat(highlight_format)
                start = pos + len(keyword)
        
        # 커서를 처음으로 이동
        cursor.setPosition(0)
        text_edit.setTextCursor(cursor)


# ============================================================================
# 빈 상태 위젯
# ============================================================================
class EmptyStateWidget(QFrame):
    """빈 상태를 표시하는 위젯"""
    
    def __init__(self, icon: str = "📂", title: str = "", description: str = ""):
        super().__init__()
        self.setObjectName("card")
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(40, 60, 40, 60)
        layout.setSpacing(15)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # 아이콘
        icon_label = QLabel(icon)
        icon_label.setStyleSheet("font-size: 48px;")
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(icon_label)
        
        # 제목
        if title:
            title_label = QLabel(title)
            title_label.setFont(QFont("", 16, QFont.Weight.Bold))
            title_label.setStyleSheet("color: #eaeaea;")
            title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(title_label)
        
        # 설명
        if description:
            desc_label = QLabel(description)
            desc_label.setStyleSheet("color: #888; font-size: 13px;")
            desc_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            desc_label.setWordWrap(True)
            layout.addWidget(desc_label)


# ============================================================================
# 프로그레스 다이얼로그
# ============================================================================
class ProgressDialog(QFrame):
    """문서 처리 진행 상황을 표시하는 다이얼로그"""
    
    canceled = pyqtSignal()
    
    def __init__(self, parent=None, title: str = "처리 중"):
        super().__init__(parent)
        self.setObjectName("card")
        self.setFixedSize(400, 180)
        
        # 그림자 효과
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(25)
        shadow.setColor(QColor(0, 0, 0, 120))
        shadow.setOffset(0, 4)
        self.setGraphicsEffect(shadow)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(25, 20, 25, 20)
        layout.setSpacing(15)
        
        # 제목
        self.title_label = QLabel(title)
        self.title_label.setFont(QFont("", 14, QFont.Weight.Bold))
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.title_label)
        
        # 현재 처리 중인 항목
        self.status_label = QLabel("준비 중...")
        self.status_label.setStyleSheet("color: #888;")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.status_label)
        
        # 프로그레스 바
        self.progress_bar = QProgressBar()
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setFormat("%p%")
        self.progress_bar.setFixedHeight(12)
        layout.addWidget(self.progress_bar)
        
        # 상세 정보
        self.detail_label = QLabel("")
        self.detail_label.setStyleSheet("color: #666; font-size: 11px;")
        self.detail_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.detail_label)
    
    def update_progress(self, percent: int, status: str):
        """진행 상황 업데이트"""
        self.progress_bar.setValue(percent)
        self.status_label.setText(status)
        self.detail_label.setText(f"{percent}% 완료")


# ============================================================================
# 메인 윈도우
# ============================================================================
class MainWindow(QMainWindow):

    def __init__(self, qa: RegulationQASystem):
        super().__init__()
        self.qa = qa
        self.history = SearchHistory()
        self.last_folder = ""
        self.model_name = AppConfig.DEFAULT_MODEL
        self.font_size = AppConfig.DEFAULT_FONT_SIZE
        self.hybrid = True
        self.worker = None
        self.status_timer = None  # 상태 레이블 타이머 관리
        
        self._load_config()
        self._init_ui()
        QTimer.singleShot(100, self._load_model)
    
    def _init_ui(self):
        self.setWindowTitle(f"{AppConfig.APP_NAME} v{AppConfig.APP_VERSION}")
        self.setMinimumSize(1000, 700)
        self.resize(1200, 800)
        
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # 헤더
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #0f3460;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(20, 0, 20, 0)
        
        logo = QLabel(f"📚 {AppConfig.APP_NAME}")
        logo.setFont(QFont("", 16, QFont.Weight.Bold))
        logo.setStyleSheet("color: white;")
        header_layout.addWidget(logo)
        header_layout.addStretch()
        
        self.status_label = QLabel("🔄 초기화 중...")
        self.status_label.setStyleSheet("color: #f59e0b;")
        header_layout.addWidget(self.status_label)
        
        version = QLabel(f"v{AppConfig.APP_VERSION}")
        version.setStyleSheet("color: #666;")
        header_layout.addWidget(version)
        
        main_layout.addWidget(header)
        
        # 탭
        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)
        main_layout.addWidget(self.tabs)
        
        self._build_search_tab()
        self._build_files_tab()
        self._build_settings_tab()
        self._setup_shortcuts()
    
    def _setup_shortcuts(self):
        """키보드 단축키 설정"""
        # Ctrl+O: 폴더 열기
        QShortcut(QKeySequence("Ctrl+O"), self).activated.connect(self._open_folder)
        # Ctrl+F: 검색창 포커스
        QShortcut(QKeySequence("Ctrl+F"), self).activated.connect(self._focus_search)
    
    def _focus_search(self):
        """검색창에 포커스"""
        self.tabs.setCurrentIndex(0)  # 검색 탭으로 이동
        self.search_input.setFocus()
        self.search_input.selectAll()
    
    def _build_search_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # 폴더 선택
        folder_frame = QFrame()
        folder_frame.setObjectName("card")
        folder_layout = QHBoxLayout(folder_frame)
        folder_layout.setContentsMargins(15, 12, 15, 12)
        
        self.folder_btn = QPushButton("📂 폴더 열기")
        self.folder_btn.setEnabled(False)
        self.folder_btn.clicked.connect(self._open_folder)
        folder_layout.addWidget(self.folder_btn)
        
        self.recent_btn = QPushButton("🕐 최근")
        self.recent_btn.setEnabled(False)
        self.recent_btn.clicked.connect(self._load_recent)
        folder_layout.addWidget(self.recent_btn)
        
        self.folder_label = QLabel("폴더를 선택하세요")
        self.folder_label.setStyleSheet("color: #888;")
        folder_layout.addWidget(self.folder_label, 1)
        
        self.refresh_btn = QPushButton("🔄")
        self.refresh_btn.setFixedWidth(40)
        self.refresh_btn.setEnabled(False)
        self.refresh_btn.clicked.connect(self._refresh)
        folder_layout.addWidget(self.refresh_btn)
        
        layout.addWidget(folder_frame)
        
        # 결과 영역
        self.result_area = QScrollArea()
        self.result_area.setWidgetResizable(True)
        self.result_container = QWidget()
        self.result_layout = QVBoxLayout(self.result_container)
        self.result_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.result_layout.setContentsMargins(10, 10, 10, 10)
        self.result_layout.setSpacing(12)
        self.result_area.setWidget(self.result_container)
        
        # 빈 상태 위젯 표시
        self._show_empty_state("welcome")
        
        layout.addWidget(self.result_area, 1)
        
        # 검색 입력
        search_frame = QFrame()
        search_frame.setObjectName("card")
        search_layout = QHBoxLayout(search_frame)
        search_layout.setContentsMargins(15, 12, 15, 12)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("검색어를 입력하세요... (최소 2글자)")
        self.search_input.setEnabled(False)
        self.search_input.returnPressed.connect(self._search)
        search_layout.addWidget(self.search_input, 1)
        
        # 검색 히스토리 버튼
        self.history_btn = QPushButton("🕑")
        self.history_btn.setFixedWidth(40)
        self.history_btn.setToolTip("최근 검색어")
        self.history_btn.clicked.connect(self._show_history_menu)
        search_layout.addWidget(self.history_btn)
        
        self.k_spin = QSpinBox()
        self.k_spin.setRange(1, 10)
        self.k_spin.setValue(AppConfig.DEFAULT_SEARCH_RESULTS)
        self.k_spin.setPrefix("결과: ")
        self.k_spin.setFixedWidth(100)
        search_layout.addWidget(self.k_spin)
        
        self.search_btn = QPushButton("🔍 검색")
        self.search_btn.setEnabled(False)
        self.search_btn.clicked.connect(self._search)
        search_layout.addWidget(self.search_btn)
        
        layout.addWidget(search_frame)
        
        self.tabs.addTab(tab, "🔍 검색")
    
    def _build_files_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 통계
        stats_frame = QFrame()
        stats_frame.setObjectName("statCard")
        stats_layout = QHBoxLayout(stats_frame)
        stats_layout.setContentsMargins(20, 15, 20, 15)
        
        self.stats_files = QLabel("📄 0개 파일")
        self.stats_files.setFont(QFont("", 12, QFont.Weight.Bold))
        stats_layout.addWidget(self.stats_files)
        
        self.stats_chunks = QLabel("📊 0 청크")
        stats_layout.addWidget(self.stats_chunks)
        
        self.stats_size = QLabel("💾 0 B")
        stats_layout.addWidget(self.stats_size)
        
        stats_layout.addStretch()
        layout.addWidget(stats_frame)
        
        # 파일 테이블
        self.file_table = QTableWidget()
        self.file_table.setColumnCount(4)
        self.file_table.setHorizontalHeaderLabels(["상태", "파일명", "크기", "청크"])
        self.file_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.file_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.file_table.setAlternatingRowColors(True)  # 교차 행 색상
        self.file_table.setSortingEnabled(True)  # 정렬 활성화
        self.file_table.setToolTip("더블클릭으로 파일 열기")
        self.file_table.doubleClicked.connect(self._open_selected_file)
        layout.addWidget(self.file_table)
        
        self.tabs.addTab(tab, "📄 파일")
    
    def _build_settings_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # 검색 설정
        search_card = QFrame()
        search_card.setObjectName("card")
        search_layout = QVBoxLayout(search_card)
        search_layout.setContentsMargins(20, 15, 20, 15)
        
        search_title = QLabel("🔍 검색 설정")
        search_title.setFont(QFont("", 13, QFont.Weight.Bold))
        search_layout.addWidget(search_title)
        
        self.hybrid_check = QCheckBox("하이브리드 검색 (벡터 + 키워드)")
        self.hybrid_check.setChecked(self.hybrid)
        self.hybrid_check.stateChanged.connect(lambda: setattr(self, 'hybrid', self.hybrid_check.isChecked()))
        self.hybrid_check.setToolTip("벡터 검색과 키워드 검색을 결합하여 더 정확한 결과 제공")
        search_layout.addWidget(self.hybrid_check)
        
        layout.addWidget(search_card)
        
        # 표시 설정
        display_card = QFrame()
        display_card.setObjectName("card")
        display_layout = QVBoxLayout(display_card)
        display_layout.setContentsMargins(20, 15, 20, 15)
        
        display_title = QLabel("🎨 표시 설정")
        display_title.setFont(QFont("", 13, QFont.Weight.Bold))
        display_layout.addWidget(display_title)
        
        # 폰트 크기 조절
        font_row = QHBoxLayout()
        font_label = QLabel("결과 폰트 크기:")
        font_row.addWidget(font_label)
        
        self.font_slider = QSlider(Qt.Orientation.Horizontal)
        self.font_slider.setRange(AppConfig.MIN_FONT_SIZE, AppConfig.MAX_FONT_SIZE)
        self.font_slider.setValue(self.font_size)
        self.font_slider.setTickInterval(2)
        self.font_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.font_slider.valueChanged.connect(self._on_font_size_changed)
        font_row.addWidget(self.font_slider, 1)
        
        self.font_size_label = QLabel(f"{self.font_size}pt")
        self.font_size_label.setFixedWidth(45)
        self.font_size_label.setStyleSheet("color: #e94560; font-weight: bold;")
        font_row.addWidget(self.font_size_label)
        
        display_layout.addLayout(font_row)
        layout.addWidget(display_card)
        
        # 모델 설정
        model_card = QFrame()
        model_card.setObjectName("card")
        model_layout = QVBoxLayout(model_card)
        model_layout.setContentsMargins(20, 15, 20, 15)
        
        model_title = QLabel("🤖 AI 모델")
        model_title.setFont(QFont("", 13, QFont.Weight.Bold))
        model_layout.addWidget(model_title)
        
        self.model_combo = QComboBox()
        self.model_combo.addItems(AppConfig.AVAILABLE_MODELS.keys())
        self.model_combo.setCurrentText(self.model_name)
        self.model_combo.currentTextChanged.connect(lambda t: setattr(self, 'model_name', t))
        model_layout.addWidget(self.model_combo)
        
        warn = QLabel("⚠️ 변경 시 재시작 필요")
        warn.setStyleSheet("color: #f59e0b; font-size: 11px;")
        model_layout.addWidget(warn)
        
        layout.addWidget(model_card)
        
        # 데이터 관리
        data_card = QFrame()
        data_card.setObjectName("card")
        data_layout = QVBoxLayout(data_card)
        data_layout.setContentsMargins(20, 15, 20, 15)
        
        data_title = QLabel("🗂️ 데이터 관리")
        data_title.setFont(QFont("", 13, QFont.Weight.Bold))
        data_layout.addWidget(data_title)
        
        btn_row = QHBoxLayout()
        clear_cache_btn = QPushButton("🗑️ 캐시 삭제")
        clear_cache_btn.setStyleSheet("background: #dc2626;")
        clear_cache_btn.clicked.connect(self._clear_cache)
        clear_cache_btn.setToolTip("벡터 인덱스 캐시를 삭제합니다")
        btn_row.addWidget(clear_cache_btn)
        
        clear_history_btn = QPushButton("🕐 히스토리 삭제")
        clear_history_btn.clicked.connect(self._clear_history)
        clear_history_btn.setToolTip("검색 히스토리를 삭제합니다")
        btn_row.addWidget(clear_history_btn)
        
        btn_row.addStretch()
        data_layout.addLayout(btn_row)
        
        layout.addWidget(data_card)
        
        # 단축키
        key_card = QFrame()
        key_card.setObjectName("card")
        key_layout = QVBoxLayout(key_card)
        key_layout.setContentsMargins(20, 15, 20, 15)
        
        key_title = QLabel("⌨️ 키보드 단축키")
        key_title.setFont(QFont("", 13, QFont.Weight.Bold))
        key_layout.addWidget(key_title)
        
        shortcuts = [("Ctrl+O", "폴더 열기"), ("Ctrl+F", "검색창 포커스"), ("Enter", "검색 실행")]
        for key, desc in shortcuts:
            row = QHBoxLayout()
            k = QLabel(key)
            k.setFont(QFont("", 11, QFont.Weight.Bold))
            k.setStyleSheet("color: #e94560;")
            k.setFixedWidth(80)
            row.addWidget(k)
            row.addWidget(QLabel(desc))
            row.addStretch()
            key_layout.addLayout(row)
        
        layout.addWidget(key_card)
        
        # 앱 정보
        info_card = QFrame()
        info_card.setObjectName("statCard")
        info_layout = QVBoxLayout(info_card)
        info_layout.setContentsMargins(20, 15, 20, 15)
        
        info_title = QLabel(f"📚 {AppConfig.APP_NAME}")
        info_title.setFont(QFont("", 14, QFont.Weight.Bold))
        info_layout.addWidget(info_title)
        
        version_label = QLabel(f"버전 {AppConfig.APP_VERSION} | PyQt6 Edition")
        version_label.setStyleSheet("color: #888;")
        info_layout.addWidget(version_label)
        
        features = QLabel("✓ 하이브리드 검색  ✓ 증분 인덱싱  ✓ 키워드 하이라이트")
        features.setStyleSheet("color: #10b981; font-size: 11px; margin-top: 5px;")
        info_layout.addWidget(features)
        
        layout.addWidget(info_card)
        layout.addStretch()
        
        self.tabs.addTab(tab, "⚙️ 설정")
    
    def _on_font_size_changed(self, value: int):
        """폰트 크기 변경 처리"""
        self.font_size = value
        self.font_size_label.setText(f"{value}pt")
        self._save_config()
    
    def _load_config(self):
        path = os.path.join(get_app_directory(), AppConfig.CONFIG_FILE)
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    cfg = json.load(f)
                self.last_folder = cfg.get("folder", "")
                self.model_name = cfg.get("model", AppConfig.DEFAULT_MODEL)
                self.font_size = cfg.get("font", AppConfig.DEFAULT_FONT_SIZE)
                self.hybrid = cfg.get("hybrid", True)
            except: pass
    
    def _save_config(self):
        try:
            with open(os.path.join(get_app_directory(), AppConfig.CONFIG_FILE), 'w', encoding='utf-8') as f:
                json.dump({"folder": self.last_folder, "model": self.model_name, "font": self.font_size, "hybrid": self.hybrid}, f)
        except: pass
    
    def _load_model(self):
        self.status_label.setText("🔄 모델 로딩 중...")
        self.worker = ModelLoaderThread(self.qa, self.model_name)
        self.worker.progress.connect(lambda m: self.status_label.setText(f"🔄 {m}"))
        self.worker.finished.connect(self._on_model_loaded)
        self.worker.start()
    
    def _on_model_loaded(self, result):
        if result.success:
            self.status_label.setText(f"✅ {result.message}")
            self.status_label.setStyleSheet("color: #10b981;")
            self.folder_btn.setEnabled(True)
            if self.last_folder and os.path.isdir(self.last_folder):
                self.recent_btn.setEnabled(True)
        else:
            self.status_label.setText(f"❌ {result.message}")
            self.status_label.setStyleSheet("color: #ef4444;")
            QMessageBox.critical(self, "오류", result.message)
    
    def _open_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "규정 폴더 선택")
        if folder:
            self._load_folder(folder)
    
    def _load_recent(self):
        if self.last_folder and os.path.isdir(self.last_folder):
            self._load_folder(self.last_folder)
    
    def _refresh(self):
        if self.last_folder:
            cache = self.qa._get_cache_dir(self.last_folder)
            shutil.rmtree(cache, ignore_errors=True)
            self._load_folder(self.last_folder)
    
    def _load_folder(self, folder):
        """폴더 로드 및 문서 처리 시작"""
        try:
            files = [os.path.join(folder, f) for f in os.listdir(folder) 
                     if os.path.splitext(f)[1].lower() in AppConfig.SUPPORTED_EXTENSIONS]
        except PermissionError:
            QMessageBox.critical(self, "오류", "폴더 접근 권한이 없습니다.")
            return
        except Exception as e:
            QMessageBox.critical(self, "오류", f"폴더 읽기 실패: {e}")
            return
        
        if not files:
            QMessageBox.warning(self, "경고", f"지원되는 파일이 없습니다.\n\n지원 형식: {', '.join(AppConfig.SUPPORTED_EXTENSIONS)}")
            return
        
        self.folder_label.setText(folder)
        self.folder_label.setToolTip(folder)
        self.folder_btn.setEnabled(False)
        
        self.progress_dialog = QMessageBox(self)
        self.progress_dialog.setWindowTitle("처리 중")
        self.progress_dialog.setText(f"문서를 처리하고 있습니다...\n({len(files)}개 파일)")
        self.progress_dialog.setStandardButtons(QMessageBox.StandardButton.NoButton)
        self.progress_dialog.show()
        
        self.worker = DocumentProcessorThread(self.qa, folder, files)
        self.worker.progress.connect(lambda p, m: self.progress_dialog.setText(f"{m}\n({p}%)"))
        self.worker.finished.connect(lambda r: self._on_folder_done(r, folder))
        self.worker.start()
    
    def _on_folder_done(self, result, folder):
        """폴더 처리 완료 핸들러"""
        self.progress_dialog.close()
        self.folder_btn.setEnabled(True)
        
        if result.success:
            self.last_folder = folder
            self._save_config()
            self.search_input.setEnabled(True)
            self.search_btn.setEnabled(True)
            self.refresh_btn.setEnabled(True)
            self.recent_btn.setEnabled(True)
            self._update_file_table()
            self._show_empty_state("ready")
            
            # 상태 표시
            self._show_status(f"✅ {result.message} (청크: {result.data.get('chunks', 0)})", "#10b981")
            self.search_input.setFocus()
            
            # 처리 실패 파일이 있으면 알림
            if result.failed_items:
                failed_count = len(result.failed_items)
                failed_list = "\n".join(result.failed_items[:5])  # 최대 5개만 표시
                more_msg = f"\n...외 {failed_count - 5}개" if failed_count > 5 else ""
                QMessageBox.warning(
                    self, 
                    "일부 파일 처리 실패",
                    f"{failed_count}개 파일 처리 실패:\n\n{failed_list}{more_msg}"
                )
        else:
            QMessageBox.critical(self, "오류", result.message)
    
    def _update_file_table(self):
        infos = self.qa.get_file_infos()
        
        # 정렬 비활성화 후 데이터 삽입 (성능 최적화)
        self.file_table.setSortingEnabled(False)
        self.file_table.setRowCount(len(infos))
        
        icons = {FileStatus.SUCCESS: "✅", FileStatus.CACHED: "💾", FileStatus.FAILED: "❌", FileStatus.PROCESSING: "⏳", FileStatus.PENDING: "⏸️"}
        total_size = 0
        total_chunks = 0
        
        for i, info in enumerate(infos):
            # 상태 아이콘
            status_item = QTableWidgetItem(icons.get(info.status, "?"))
            status_item.setData(Qt.ItemDataRole.UserRole, info.path)  # 파일 경로 저장
            self.file_table.setItem(i, 0, status_item)
            
            # 파일명 (경로 저장)
            name_item = QTableWidgetItem(info.name)
            name_item.setData(Qt.ItemDataRole.UserRole, info.path)
            name_item.setToolTip(info.path)  # 전체 경로 툴팁
            self.file_table.setItem(i, 1, name_item)
            
            # 크기
            size_item = QTableWidgetItem(FileUtils.format_size(info.size))
            size_item.setData(Qt.ItemDataRole.UserRole + 1, info.size)  # 정렬용 숫자 저장
            self.file_table.setItem(i, 2, size_item)
            
            # 청크
            chunk_item = QTableWidgetItem(str(info.chunks))
            chunk_item.setData(Qt.ItemDataRole.UserRole + 1, info.chunks)  # 정렬용 숫자 저장
            self.file_table.setItem(i, 3, chunk_item)
            
            total_size += info.size
            total_chunks += info.chunks
        
        # 정렬 다시 활성화
        self.file_table.setSortingEnabled(True)
        
        self.stats_files.setText(f"📄 {len(infos)}개 파일")
        self.stats_chunks.setText(f"📊 {total_chunks} 청크")
        self.stats_size.setText(f"💾 {FileUtils.format_size(total_size)}")
    
    def _open_selected_file(self):
        """선택된 파일 열기 (정렬과 무관하게 작동)"""
        row = self.file_table.currentRow()
        if row >= 0:
            # 저장된 파일 경로 가져오기
            name_item = self.file_table.item(row, 1)
            if name_item:
                file_path = name_item.data(Qt.ItemDataRole.UserRole)
                if file_path:
                    FileUtils.open_file(file_path)
    
    def _search(self):
        query = self.search_input.text().strip()
        if not query:
            return
        if not self.qa.vector_store:
            QMessageBox.warning(self, "경고", "문서를 먼저 로드하세요")
            return
        
        self.search_btn.setEnabled(False)
        self._clear_results()
        loading = QLabel("🔍 검색 중...")
        loading.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.result_layout.addWidget(loading)
        
        # 검색 시간 측정 시작
        import time
        self._search_start_time = time.time()
        
        self.worker = SearchThread(self.qa, query, self.k_spin.value(), self.hybrid)
        self.worker.finished.connect(lambda r: self._on_search_done(r, query))
        self.worker.start()
    
    def _on_search_done(self, result, query):
        import time
        search_time = time.time() - getattr(self, '_search_start_time', time.time())
        
        self.search_btn.setEnabled(True)
        self._clear_results()
        
        if not result.success:
            err = QLabel(f"❌ {result.message}")
            err.setStyleSheet("color: #ef4444;")
            err.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.result_layout.addWidget(err)
            return
        
        if not result.data:
            self._show_empty_state("no_results")
            return
        
        self.history.add(query)
        self.last_search_results = result.data  # 내보내기용 저장
        self.last_search_query = query
        
        # 결과 헤더 (검색어 + 통계 + 내보내기 버튼)
        header_frame = QFrame()
        header_frame.setObjectName("card")
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(15, 10, 15, 10)
        
        query_label = QLabel(f"🔎 \"{query}\" - {len(result.data)}개 결과")
        query_label.setFont(QFont("", 12, QFont.Weight.Bold))
        header_layout.addWidget(query_label)
        
        # 검색 시간 표시
        time_label = QLabel(f"⏱ {search_time:.2f}초")
        time_label.setStyleSheet("color: #888; font-size: 11px;")
        header_layout.addWidget(time_label)
        
        header_layout.addStretch()
        
        # 내보내기 버튼
        export_btn = QPushButton("📥 내보내기")
        export_btn.setFixedHeight(30)
        export_btn.clicked.connect(self._export_results)
        header_layout.addWidget(export_btn)
        
        self.result_layout.addWidget(header_frame)
        
        for i, item in enumerate(result.data, 1):
            card = ResultCard(i, item, self._copy_text, self.font_size, query)
            self.result_layout.addWidget(card)
        
        self.search_input.clear()
        self.search_input.setFocus()
    
    def _export_results(self):
        """검색 결과 내보내기"""
        if not hasattr(self, 'last_search_results') or not self.last_search_results:
            QMessageBox.warning(self, "알림", "내보낼 검색 결과가 없습니다.")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "검색 결과 내보내기",
            f"검색결과_{self.last_search_query[:20]}.txt",
            "텍스트 파일 (*.txt);;CSV 파일 (*.csv)"
        )
        
        if not file_path:
            return
        
        try:
            is_csv = file_path.lower().endswith('.csv')
            
            with open(file_path, 'w', encoding='utf-8') as f:
                if is_csv:
                    f.write("순위,점수,파일,내용\n")
                    for i, item in enumerate(self.last_search_results, 1):
                        content = item['content'].replace('"', '""').replace('\n', ' ')
                        f.write(f'{i},{item["score"]:.2f},"{item["source"]}","{content}"\n')
                else:
                    f.write(f"검색어: {self.last_search_query}\n")
                    f.write(f"결과 수: {len(self.last_search_results)}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    for i, item in enumerate(self.last_search_results, 1):
                        f.write(f"[결과 {i}] ({int(item['score']*100)}%)\n")
                        f.write(f"파일: {item['source']}\n")
                        f.write("-" * 30 + "\n")
                        f.write(item['content'] + "\n\n")
            
            self._show_status(f"✅ 결과 내보내기 완료: {os.path.basename(file_path)}", "#10b981", 3000)
        except Exception as e:
            QMessageBox.critical(self, "오류", f"내보내기 실패: {e}")
    
    def _clear_results(self):
        while self.result_layout.count():
            item = self.result_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
    
    def _copy_text(self, text):
        """텍스트 복사 및 상태 표시"""
        QApplication.clipboard().setText(text)
        self._show_status("✅ 클립보드에 복사됨", "#10b981", 2000)
    
    def _show_status(self, message: str, color: str = "#eaeaea", duration: int = 0):
        """상태 레이블에 메시지 표시 (duration이 0이면 영구 표시)"""
        # 이전 타이머 취소
        if self.status_timer:
            self.status_timer.stop()
            self.status_timer = None
        
        self.status_label.setText(message)
        self.status_label.setStyleSheet(f"color: {color};")
        
        if duration > 0:
            self.status_timer = QTimer()
            self.status_timer.setSingleShot(True)
            self.status_timer.timeout.connect(lambda: self.status_label.setText(""))
            self.status_timer.start(duration)
    
    def _show_empty_state(self, state_type: str = "welcome"):
        """빈 상태 위젯 표시"""
        self._clear_results()
        
        if state_type == "welcome":
            widget = EmptyStateWidget(
                "👋",
                "사내 규정 검색기",
                "폴더를 선택하고 문서를 로드한 후 검색을 시작하세요.\nCtrl+O로 폴더 열기"
            )
        elif state_type == "no_results":
            widget = EmptyStateWidget(
                "🔍",
                "검색 결과 없음",
                "다른 검색어로 시도해보세요."
            )
        elif state_type == "ready":
            widget = EmptyStateWidget(
                "✅",
                "검색 준비 완료",
                "검색어를 입력하고 Enter를 누르거나 검색 버튼을 클릭하세요."
            )
        else:
            return
        
        self.result_layout.addWidget(widget)
    
    def _show_history_menu(self):
        """검색 히스토리 메뉴 표시"""
        history_items = self.history.get(10)
        
        if not history_items:
            QMessageBox.information(self, "알림", "검색 히스토리가 없습니다.")
            return
        
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background: #16213e;
                border: 1px solid #e94560;
                border-radius: 6px;
                padding: 5px;
            }
            QMenu::item {
                background: transparent;
                color: white;
                padding: 8px 20px;
                border-radius: 4px;
            }
            QMenu::item:selected {
                background: #e94560;
            }
        """)
        
        for query in history_items:
            action = menu.addAction(f"🔍 {query}")
            action.triggered.connect(lambda checked, q=query: self._search_from_history(q))
        
        menu.addSeparator()
        clear_action = menu.addAction("🗑️ 히스토리 삭제")
        clear_action.triggered.connect(self._clear_history)
        
        # 버튼 아래에 메뉴 표시
        menu.exec(self.history_btn.mapToGlobal(self.history_btn.rect().bottomLeft()))
    
    def _search_from_history(self, query: str):
        """히스토리에서 선택한 검색어로 검색"""
        self.search_input.setText(query)
        self._search()
    
    def _clear_cache(self):
        if QMessageBox.question(self, "확인", "캐시를 삭제하시겠습니까?") == QMessageBox.StandardButton.Yes:
            self.qa.clear_cache()
            self._show_status("✅ 캐시 삭제됨", "#10b981", 3000)
    
    def _clear_history(self):
        if QMessageBox.question(self, "확인", "히스토리를 삭제하시겠습니까?") == QMessageBox.StandardButton.Yes:
            self.history.clear()
            self._show_status("✅ 히스토리 삭제됨", "#10b981", 3000)
    
    def closeEvent(self, event):
        self._save_config()
        self.qa.cleanup()
        event.accept()


# ============================================================================
# 메인
# ============================================================================
def main():
    try:
        app = QApplication(sys.argv)
        app.setStyle('Fusion')
        app.setStyleSheet(DARK_STYLE)
        
        qa = RegulationQASystem()
        window = MainWindow(qa)
        window.show()
        
        sys.exit(app.exec())
    except Exception as e:
        logger.critical(f"치명적 오류: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
